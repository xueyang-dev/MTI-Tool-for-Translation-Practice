"""Literature Evidence Spine tests.

Run: .venv/bin/python tests/literature_evidence_spine_test.py
"""
from __future__ import annotations

import json
import shutil
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document

from mti_tool import academic_evidence, academic_validator, academic_writer
from mti_tool import literature_evidence
from tests.academic_writing_test import JOB, _state


def _registered_source(content: str = "# Theory\nFunctional equivalence focuses on receptor response."):
    return {
        "source_id": "nida1964",
        "title": "Toward a Science of Translating",
        "authors": ["Eugene Nida"],
        "year": 1964,
        "source_type": "book",
        "verification_status": "metadata_verified",
        "allowed_citation_status": "allowed",
        "citation_metadata": {
            "title": "Toward a Science of Translating",
            "authors": ["Eugene Nida"],
            "year": 1964,
            "in_text": "Nida（1964）",
        },
        "content": content,
        "content_format": "markdown",
        "notes": [{"note_id": "N-1", "text": "用于解释受众反应。"}],
        "manual_excerpts": [{
            "excerpt_id": "EX-1",
            "text": "Translation consists in reproducing the closest natural equivalent.",
            "location": {"page": 12},
        }],
        "extracted_passages": [{
            "passage_id": "MX-1", "text": "A model-extracted passage requiring review.",
            "location": {"page": 13},
        }],
    }


class ClaimMock:
    def __call__(self, provider, api_key, model, system_prompt, user_prompt,
                 temperature=0.1):
        if "文献主张抽取器" not in system_prompt:
            return "{}"
        payload = json.loads(user_prompt)
        source_passage = next(
            x for x in payload["evidence"]
            if x["provenance"] == "source_text_verified")
        return json.dumps({"claims": [
            {
                "statement": "功能对等关注接受者反应。",
                "source_id": source_passage["source_id"],
                "evidence_ids": [source_passage["evidence_id"]],
                "claim_type": "theoretical_position",
                "confidence": "medium",
            },
            {
                "statement": "无效主张",
                "source_id": "unknown",
                "evidence_ids": ["LE-unknown"],
                "claim_type": "empirical_finding",
                "confidence": "high",
            },
        ]}, ensure_ascii=False)


def _literature_artifacts(content: str | None = None):
    source_artifact = literature_evidence.build_literature_sources([
        _registered_source(content or "# Theory\nFunctional equivalence focuses on receptor response."),
        {
            "source_id": "metadata-only", "title": "Metadata Only",
            "authors": ["A. Author"], "year": 2020,
            "verification_status": "candidate",
        },
    ])
    evidence_artifact = literature_evidence.build_literature_evidence(source_artifact)
    claims_artifact = literature_evidence.build_literature_claims(
        source_artifact, evidence_artifact, ClaimMock(), "x", "x", "x")
    return source_artifact, evidence_artifact, claims_artifact


def test_source_formats_locations_and_hashes():
    tmp = Path(tempfile.mkdtemp(prefix="lit-source-formats-"))
    try:
        markdown = tmp / "paper.md"
        markdown.write_text("# A\nFirst paragraph.\n\n## B\nSecond paragraph.", encoding="utf-8")
        docx = tmp / "paper.docx"
        document = Document()
        document.add_heading("Heading", level=1)
        document.add_paragraph("A DOCX paragraph.")
        document.save(docx)

        import fitz
        pdf = tmp / "paper.pdf"
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        page.insert_text((72, 72), "A PDF page passage.")
        pdf_doc.save(pdf)
        pdf_doc.close()

        artifact = literature_evidence.build_literature_sources([
            {**_registered_source(), "source_id": "md", "local_source_path": str(markdown),
             "content": None, "notes": [], "manual_excerpts": [], "extracted_passages": []},
            {**_registered_source(), "source_id": "docx", "local_source_path": str(docx),
             "content": None, "notes": [], "manual_excerpts": [], "extracted_passages": []},
            {**_registered_source(), "source_id": "pdf", "local_source_path": str(pdf),
             "content": None, "notes": [], "manual_excerpts": [], "extracted_passages": []},
        ])
        by_id = literature_evidence.source_index(artifact)
        assert any(x["location"]["kind"] == "markdown" for x in by_id["md"]["content_blocks"])
        assert any(x["location"]["kind"] == "docx_paragraph"
                   for x in by_id["docx"]["content_blocks"])
        assert any(x["location"]["kind"] == "pdf_page"
                   for x in by_id["pdf"]["content_blocks"])
        old_hash = by_id["md"]["content_hash"]
        markdown.write_text("# A\nChanged paragraph.", encoding="utf-8")
        changed = literature_evidence.build_literature_sources([
            {**_registered_source(), "source_id": "md", "local_source_path": str(markdown),
             "content": None, "notes": [], "manual_excerpts": [], "extracted_passages": []},
        ])
        assert literature_evidence.source_index(changed)["md"]["content_hash"] != old_hash
        print("  ✓ PDF/Markdown/DOCX 精确位置与来源内容哈希")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_evidence_origins_claims_and_metadata_only():
    sources, evidence, claims = _literature_artifacts()
    source = literature_evidence.source_index(sources)["nida1964"]
    assert source["verification_status"] == "metadata_verified"
    assert source["allowed_citation_status"] == "allowed"
    assert source["content_availability"] == "full_text_available"
    items = [x for x in evidence["items"] if x["source_id"] == "nida1964"]
    assert {x["provenance"] for x in items} == {
        "source_text_verified", "user_note", "manual_excerpt",
        "model_extracted_from_source",
    }
    metadata = next(x for x in evidence["items"] if x["source_id"] == "metadata-only")
    assert metadata["provenance"] == "metadata_only" and not metadata["eligible_for_claim"]
    assert len(claims["items"]) == 1
    claim = claims["items"][0]
    assert claim["literature_claim_id"] == "LC-001"
    assert claim["evidence_grounded_status"] == "grounded"
    assert claim["supporting_evidence_ids"] != ["LE-unknown"]
    print("  ✓ 五类 provenance、metadata-only 降级与 Literature Claim 关系")


def _validation_graph():
    sources, lit_evidence, lit_claims = _literature_artifacts()
    project = academic_evidence.build_academic_evidence(_state(), JOB)
    research = academic_writer.build_research_model(
        project, "功能对等", {"research_questions": ["如何解释受众反应？"]})
    lc = lit_claims["items"][0]
    le = lc["supporting_evidence_ids"][0]
    argument = {"claims": [{
        "claim_id": "C1", "claim": "该案例可从受众反应角度作有限解释。",
        "research_question": "RQ1", "project_evidence": [],
        "literature_claims": [lc["literature_claim_id"]],
        "literature_evidence": [le], "support_category": "literature_supported",
        "analysis_type": "AUTHOR_ANALYSIS", "confidence": "medium",
    }]}
    outline = {"sections": [{
        "section_id": "1", "title": "文献论证", "research_questions": ["RQ1"],
        "claims": ["C1"], "cases": [],
        "literature_claims": [lc["literature_claim_id"]],
        "literature_evidence": [le], "literature_sources": ["nida1964"],
        "required_statistics": [], "minimum_chars": 20,
    }]}
    selected = {"cases": []}
    exact = literature_evidence.evidence_index(lit_evidence)[le]["evidence_text"]
    report = (
        "## 1 文献论证\n\n<!--rq:RQ1--><!--claim:C1-->"
        f"<!--lit-claim:{lc['literature_claim_id']}--><!--lit-evidence:{le}-->"
        "从结果看，该理论可用于解释本项目，但这属于作者分析 [@nida1964]。\n"
        f"> [LITERATURE {le}]: {exact}\n"
        "该结论不超出所列文献证据。"
    )
    report = academic_validator.expand_evidence_tokens(
        report, {**project, "literature_sources": sources["sources"]})
    return project, research, argument, selected, outline, sources, lit_evidence, lit_claims, report


def test_validator_full_literature_integrity():
    graph = _validation_graph()
    valid = academic_validator.validate_academic_report(
        graph[8], graph[0], graph[1], graph[2], graph[3], graph[4],
        graph[5], graph[6], graph[7])
    assert valid["status"] == "pass", valid["issues"]

    project, research, argument, selected, outline, sources, evidence, claims, report = graph
    tampered_sources = json.loads(json.dumps(sources))
    tampered_sources["sources"][0]["content_hash"] = "tampered"
    tampered_evidence = json.loads(json.dumps(evidence))
    usable = next(x for x in tampered_evidence["items"] if x.get("eligible_for_claim"))
    usable["location"] = {"kind": "pdf_page", "page": 999}
    usable["evidence_text"] = "fabricated quote"
    bad_argument = json.loads(json.dumps(argument))
    bad_argument["claims"][0]["literature_evidence"] = ["nida1964", "LE-unknown"]
    bad_outline = json.loads(json.dumps(outline))
    bad_outline["sections"][0]["literature_sources"] = []
    bad_report = report.replace(
        "该结论不超出所列文献证据。",
        "该结论不超出证据。 Wrong (2020) [@evil] [@metadata-only]。")
    result = academic_validator.validate_academic_report(
        bad_report, project, research, bad_argument, selected, bad_outline,
        tampered_sources, tampered_evidence, claims)
    issue_types = {x["type"] for x in result["issues"]}
    assert {
        "literature_source_hash_mismatch", "invalid_literature_location",
        "literature_evidence_text_mismatch", "literature_evidence_hash_mismatch",
        "argument_plan_source_id_without_grounding",
        "global_claim_unknown_literature_evidence", "unknown_literature_citation",
        "uncitable_literature_source", "section_literature_outside_plan",
    }.issubset(issue_types), issue_types
    print("  ✓ 来源/位置/逐字文本/哈希/主张关系/未登记引用确定性校验")


class LiteraturePipelineMock:
    def __init__(self):
        self.inject_unregistered = True
        self.repaired_sections = []
        self.literature_review_calls = 0

    def __call__(self, provider, api_key, model, system_prompt, user_prompt,
                 temperature=0.1):
        if "文献主张抽取器" in system_prompt:
            payload = json.loads(user_prompt)
            item = next(x for x in payload["evidence"]
                        if x["provenance"] == "source_text_verified")
            return json.dumps({"claims": [{
                "statement": "功能对等关注接受者反应。",
                "source_id": item["source_id"], "evidence_ids": [item["evidence_id"]],
                "claim_type": "theoretical_position", "confidence": "medium",
            }]}, ensure_ascii=False)
        if "学术论证规划器" in system_prompt:
            payload = json.loads(user_prompt)
            lc = payload["literature_claims"][0]
            return json.dumps({"claims": [{
                "claim": "翻译决策可从接受者反应角度作有限解释。",
                "research_question": "RQ1", "project_evidence": [],
                "literature_claims": [lc["literature_claim_id"]],
                "literature_evidence": lc["supporting_evidence_ids"],
                "support_category": "literature_supported",
                "analysis_type": "AUTHOR_ANALYSIS", "confidence": "medium",
                "planned_sections": ["2"], "reasoning": "文献证据已落地。",
                "counterargument": "不能把理论解释当成译者心理事实。",
            }]}, ensure_ascii=False)
        if "学术提纲规划器" in system_prompt:
            payload = json.loads(user_prompt)
            claim = payload["argument_plan"]["claims"][0]
            sections = []
            for sid, title in (("1", "研究设计"), ("2", "理论分析"), ("3", "结论")):
                has_claim = sid == "2"
                sections.append({
                    "section_id": sid, "title": title, "purpose": title,
                    "research_questions": ["RQ1"],
                    "claims": ["C1"] if has_claim else [], "cases": [],
                    "literature_claims": claim["literature_claims"] if has_claim else [],
                    "literature_evidence": claim["literature_evidence"] if has_claim else [],
                    "required_statistics": [], "target_words": 200, "minimum_chars": 60,
                    "allowed_conclusions": ["不超过证据范围"],
                })
            return json.dumps({"sections": sections}, ensure_ascii=False)
        if "独立的 MTI 学术审稿人" in system_prompt:
            return '{"issues":[]}'
        if "Literature Support Reviewer" in system_prompt:
            self.literature_review_calls += 1
            if self.literature_review_calls == 1:
                payload = json.loads(user_prompt)
                global_claim = payload["argument_plan"]["claims"][0]
                literature_claim = payload["literature_claims"][0]
                return json.dumps({"issues": [{
                    "type": "claim_too_broad", "section_id": "2",
                    "global_claim_id": global_claim["claim_id"],
                    "literature_claim_id": literature_claim["literature_claim_id"],
                    "literature_evidence_ids": literature_claim[
                        "supporting_evidence_ids"],
                    "source_id": literature_claim["source_id"],
                    "severity": "medium",
                    "reason": "全局表述强于所给文献段落。",
                    "repair_action": "narrow",
                }]}, ensure_ascii=False)
            return '{"issues":[]}'
        if "证据约束型学术写作者" in system_prompt:
            payload = json.loads(user_prompt)
            packet = payload["packet"]
            section = packet["current_section"]
            if "existing_section" in payload:
                self.repaired_sections.append(section["section_id"])
            body = "".join(f"<!--rq:{x}-->" for x in section["research_questions"])
            body += "".join(f"<!--claim:{x}-->" for x in section["claims"])
            for claim in packet["literature_claims"]:
                body += f"<!--lit-claim:{claim['literature_claim_id']}-->"
            for evidence in packet["literature_evidence"]:
                body += f"<!--lit-evidence:{evidence['evidence_id']}-->"
                body += f"文献主张由逐字证据限定 [@{evidence['source_id']}]。"
                body += (f"\n> [LITERATURE {evidence['evidence_id']}]: "
                         f"{evidence['evidence_text']}\n")
            body += "本节只在已登记证据范围内展开作者分析，不还原不可观察的心理意图。" * 6
            if section["section_id"] == "2" and "existing_section" not in payload \
                    and self.inject_unregistered:
                body += " 虚构来源 [@evil2026]。"
            return body
        return "{}"


def test_end_to_end_trace_repair_and_precise_staleness():
    tmp = Path(tempfile.mkdtemp(prefix="literature-spine-e2e-"))
    try:
        state = _state()
        mock = LiteraturePipelineMock()
        literature = [_registered_source()]
        academic_writer.run_academic_pipeline(
            state, JOB, "功能对等", "x", "x", "x", tmp,
            mock, lambda current: None,
            research_settings={"research_questions": ["如何解释受众反应？"]},
            literature_sources=literature, auto_repair_rounds=1)
        assert state["p3_done"]
        assert mock.repaired_sections == ["2"]
        # Canonical repair artifact carries Literature Claim -> Global Claim -> section.
        repair_artifact = academic_writer._read_artifact(tmp / "academic-repair-history.json")
        assert repair_artifact["rounds"][0]["literature_claim_ids"] == ["LC-001"]
        assert repair_artifact["rounds"][0]["global_claim_ids"] == ["C1"]
        assert repair_artifact["rounds"][0]["repair_actions"] == ["narrow"]
        assert state["academic_state"]["validation_history"][0]["status"] == "fail"
        assert state["academic_state"]["validation_history"][-1]["status"] == "pass"
        assert "evil2026" not in state["p3_md"]
        expected = {
            "literature-sources.json", "literature-evidence.jsonl",
            "literature-claims.jsonl", "literature-support-review.json",
            "argument-plan.json", "academic-outline.json", "academic-sections.json",
        }
        assert expected.issubset({x.name for x in tmp.iterdir()})
        sources = academic_writer._read_artifact(tmp / "literature-sources.json")
        evidence = academic_writer._read_artifact(tmp / "literature-evidence.jsonl")
        claims = academic_writer._read_artifact(tmp / "literature-claims.jsonl")
        argument = academic_writer._read_artifact(tmp / "argument-plan.json")
        outline = academic_writer._read_artifact(tmp / "academic-outline.json")
        sections = academic_writer._read_artifact(tmp / "academic-sections.json")
        lc = claims["items"][0]
        le = lc["supporting_evidence_ids"][0]
        global_claim = argument["claims"][0]
        section = next(x for x in outline["sections"] if "C1" in x["claims"])
        section_artifact = next(x for x in sections["sections"]
                                if x["section_id"] == section["section_id"])
        assert global_claim["literature_claims"] == [lc["literature_claim_id"]]
        assert global_claim["literature_evidence"] == [le]
        assert le in section["literature_evidence"]
        assert le in section_artifact["provenance"]["literature_evidence_ids"]
        assert literature_evidence.evidence_index(evidence)[le]["source_id"] == \
            literature_evidence.source_index(sources)["nida1964"]["source_id"]

        dependencies = {
            name: state["academic_state"]["artifacts"][name]["dependency_hash"]
            for name in ("evidence", "literature_sources", "literature_evidence",
                         "literature_claims", "argument_plan", "sections")}
        mock.inject_unregistered = False
        metadata_changed = [{**literature[0], "citation_metadata": {
            **literature[0]["citation_metadata"], "in_text": "Nida 1964"}}]
        academic_writer.run_academic_pipeline(
            state, JOB, "功能对等", "x", "x", "x", tmp,
            mock, lambda current: None,
            research_settings={"research_questions": ["如何解释受众反应？"]},
            literature_sources=metadata_changed, auto_repair_rounds=0)
        after_metadata = state["academic_state"]["artifacts"]
        assert after_metadata["evidence"]["dependency_hash"] == dependencies["evidence"]
        assert after_metadata["literature_evidence"]["dependency_hash"] == \
            dependencies["literature_evidence"]
        assert after_metadata["literature_claims"]["dependency_hash"] == \
            dependencies["literature_claims"]
        assert after_metadata["literature_sources"]["dependency_hash"] != \
            dependencies["literature_sources"]

        before_content = {
            name: after_metadata[name]["dependency_hash"]
            for name in ("evidence", "literature_evidence", "literature_claims",
                         "argument_plan", "sections")}
        content_changed = [{**metadata_changed[0], "content":
                            "# Theory\nChanged source passage with narrower meaning."}]
        academic_writer.run_academic_pipeline(
            state, JOB, "功能对等", "x", "x", "x", tmp,
            mock, lambda current: None,
            research_settings={"research_questions": ["如何解释受众反应？"]},
            literature_sources=content_changed, auto_repair_rounds=0)
        after_content = state["academic_state"]["artifacts"]
        assert after_content["evidence"]["dependency_hash"] == before_content["evidence"]
        for name in ("literature_evidence", "literature_claims", "argument_plan", "sections"):
            assert after_content[name]["dependency_hash"] != before_content[name], name

        versions = dict(academic_writer.VERSIONS,
                        writer_version="writer-version-staleness-test")
        academic_writer.sync_versions(state, versions)
        assert "literature_sources" in state["academic_state"]["artifacts"]
        assert "literature_evidence" in state["academic_state"]["artifacts"]
        assert "literature_claims" in state["academic_state"]["artifacts"]
        assert "sections" not in state["academic_state"]["artifacts"]
        print("  ✓ E2E trace、恶意未登记引用定点修复与精确失效")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_historical_metadata_only_downgrade():
    sources = literature_evidence.build_literature_sources([{
        "source_id": "legacy-paper", "title": "Legacy Paper",
        "authors": ["Legacy Author"], "year": 2001,
        "source_status": "unverified_candidate", "citation_allowed": True,
    }])
    source = sources["sources"][0]
    evidence = literature_evidence.build_literature_evidence(sources)
    assert source["verification_status"] == "candidate"
    assert source["allowed_citation_status"] == "not_allowed"
    assert source["content_availability"] == "metadata_only"
    assert evidence["items"][0]["verification_status"] == "metadata_only"
    assert not evidence["items"][0]["eligible_for_claim"]
    print("  ✓ 历史 paper metadata 保留但降级为 evidence_missing")


if __name__ == "__main__":
    print("文献证据脊柱测试：")
    test_source_formats_locations_and_hashes()
    test_evidence_origins_claims_and_metadata_only()
    test_validator_full_literature_integrity()
    test_end_to_end_trace_repair_and_precise_staleness()
    test_historical_metadata_only_downgrade()
    print("\n全部通过 ✅")
