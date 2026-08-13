"""Regression checks for the college-formatted thesis review document."""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "eval/academic-quality/ec100d8686d3891e/thesis-closeout-v6"


def main() -> None:
    docx_path = OUT / "thesis-review-v6.docx"
    manifest = json.loads((OUT / "thesis-document-assembly-manifest.json").read_text())
    closeout = json.loads((OUT / "thesis-closeout-state.json").read_text())
    assert docx_path.is_file() and docx_path.stat().st_size > 80_000
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    for required in (
        "摘  要", "ABSTRACT", "目  录", "1 引言", "2 翻译项目概述",
        "3 翻译项目案例分析", "4 总结与反思", "参考文献", "致  谢",
        "附录一 核心案例原文、初译与终译", "附录二 翻译实践原文与终译",
    ):
        assert required in text
    assert len(doc.sections) == 5
    assert len(doc.tables) >= 5
    assert manifest["appendix_translation_pairs"] == 273
    assert manifest["supervisor_review_status"] == "pending_supervisor_review"
    assert closeout["document_assembly"]["status"] == "review_draft_assembled"
    assert closeout["phase_b"]["semantic_review_status"] == "pending_supervisor_review"
    assert "\t—" not in text
    with zipfile.ZipFile(docx_path) as archive:
        xml_text = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist() if name.endswith(".xml")
        )
    assert not re.search(r"seg-ec100d8686d3891e|<!--(?:claim|rq|stat|cite):", xml_text)
    print("thesis document assembly: PASS")


if __name__ == "__main__":
    main()
