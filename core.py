"""MTI 翻译实践小助手 —— 核心逻辑层（与 Streamlit UI 解耦，便于测试）。

职责：大模型路由、文档清洗、术语抽取、双语翻译、报告生成、任务进度持久化。
"""
import hashlib
import io
import json
import re
import shutil
import time
from datetime import datetime, timezone
from pathlib import Path

import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from google import genai
from openai import OpenAI

from mti_tool import models as _models
from mti_tool import state_migration as _state_migration

# ================= 常量 =================
# 任务进度与过程文件的本地存储目录（已加入 .gitignore）
OUTPUT_DIR = Path("outputs")

# 各家模型可选列表（UI 中可切换；如模型下线，在这里更换即可）
MODELS = {
    "DeepSeek": ["deepseek-chat", "deepseek-reasoner"],
    "OpenAI": ["gpt-4o-mini", "gpt-4o", "gpt-4.1-mini", "gpt-4.1"],
    "Gemini": ["gemini-2.5-flash", "gemini-2.5-pro", "gemini-2.0-flash"],
}


# ================= 基础工具函数 =================
def clean_xml_chars(text):
    if not isinstance(text, str):
        return str(text)
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)


def parse_json_array(text):
    """从 LLM 输出中稳健地解析 JSON 数组；解析失败返回 None。

    依次尝试：整体解析 -> 去掉 Markdown 代码块后解析 -> 从每个 '[' 位置做
    raw_decode（可容忍输出前后夹带解释文字）。
    """
    if not isinstance(text, str) or not text.strip():
        return None
    candidate = text.strip()
    candidate = re.sub(r'^```(?:json)?\s*', '', candidate, flags=re.DOTALL)
    candidate = re.sub(r'\s*```$', '', candidate, flags=re.DOTALL).strip()

    try:
        obj = json.loads(candidate)
        if isinstance(obj, list):
            return obj
    except Exception:
        pass

    decoder = json.JSONDecoder()
    for m in re.finditer(r'\[', candidate):
        try:
            obj, _ = decoder.raw_decode(candidate[m.start():])
        except Exception:
            continue
        if isinstance(obj, list):
            return obj
    return None


def _translation_item_text(item):
    """把模型返回数组项归一为译文文本：支持字符串或 {translation/target/...} 对象。"""
    if isinstance(item, str):
        return item
    if isinstance(item, dict):
        for key in ("translation", "target", "text", "译文", "翻译", "content"):
            val = item.get(key)
            if isinstance(val, str) and val.strip():
                return val
        for val in item.values():
            if isinstance(val, str) and val.strip():
                return val
    return ""


def parse_translation_array(res, expected):
    """解析翻译响应，支持多种模型实际输出形态：
    1. JSON 数组；2. JSON 对象 {"1": "..."}；3. 编号行（1. 译文 / 1、译文）；
    4. 单段时接受任意非空文本。全部失败返回 None。
    """
    arr = parse_json_array(res)
    if arr is not None and len(arr) == expected:
        out = [_translation_item_text(item) for item in arr]
        if all(t.strip() for t in out):
            return out
    if isinstance(res, str):
        candidate = res.strip()
        candidate = re.sub(r'^```(?:json)?\s*', '', candidate, flags=re.DOTALL)
        candidate = re.sub(r'\s*```$', '', candidate, flags=re.DOTALL).strip()
        try:
            obj = json.loads(candidate)
        except Exception:
            obj = None
        if isinstance(obj, dict):
            out = []
            for i in range(1, expected + 1):
                v = obj.get(str(i))
                if not isinstance(v, str) or not v.strip():
                    return None
                out.append(v.strip())
            return out
    if isinstance(res, str):
        numbered = re.findall(r'^\s*(\d+)[.)、]\s*(.+?)\s*$', res, re.M)
        if numbered:
            numbered.sort(key=lambda x: int(x[0]))
            texts = [t for _, t in numbered]
            if len(texts) >= expected:
                out = [t for t in texts[:expected] if t.strip()]
                if len(out) == expected:
                    return out
    if expected == 1 and isinstance(res, str) and res.strip():
        # 单段兜底：去掉编号前缀后按整段接受，交给确定性检查与审校把关
        raw = res.strip()
        if raw.startswith('[') or raw.startswith('{'):
            return None  # 形似 JSON 的响应不允许当纯文本吞下
        return [re.sub(r'^\s*\d+[.)、]\s*', '', raw).strip()]
    return None


def is_rate_limited(err):
    s = str(err)
    return '429' in s or 'RESOURCE_EXHAUSTED' in s or 'rate limit' in s.lower()


# ================= PDF 确定性段落提取 =================
# 经验（来自 localize-anything 与全书实测）：分段/清洗必须确定性，不能交给 LLM。
# 旧流程把 ~2500 字符的任意文本块交给模型"清洗"，导致两类系统性缺陷：
#   1. 块边界落在句中 -> 句子被拦腰截断（"…pecking at" / "crumbs and bones…"）；
#   2. 模型自由裁量 -> 对白、引语被随意拆分或合并，分段结果不可复现。
# 新流程直接读 PDF 版面：块(block)->行(line)->首行缩进判定段落，连字符修复、
# 跨页段落合并、页眉页脚/页码剔除全部确定完成。

# 句末终结符（用于判断段落是否未完结、需要与下一段合并）
_SENTENCE_TERMINAL = set('.!?"”’…:;)')

# 纯装饰符号行（章节分隔花饰等），无字母/数字，不是正文
_ORNAMENT_RE = re.compile(r"^[\s*•·▪◦‣❦❧—–\-]{1,12}$")

# 常见缩写（句点不计入句界）
_ABBREV_RE = re.compile(
    r"\b(?:Lt|Col|Gen|Maj|Capt|Sgt|Brig|Mr|Mrs|Ms|Dr|St|No|Vol|pp|"
    r"e\.g|i\.e|vs|etc|a\.m|p\.m|U\.S|A\.F|B\.C|A\.D)\.", re.IGNORECASE)


def extract_pdf_paragraphs(file_bytes):
    """从 PDF 确定性重建段落列表。

    规则：
    - 行内文本按 span 拼接；行 -> 段落依据首行缩进（x0 明显大于正文 x0 即新段落）；
    - 连字符换行修复（"word-" + 小写开头 -> 去连字符直接拼接）；
    - 跨页/跨块延续：上一段未以终结符结尾且长度超过标题阈值 -> 合并；
    - 小写开头的碎片段 -> 并入上一段（碎句兜底）；
    - 剔除页眉页脚（跨 ≥20% 页重复出现的短行）与独立页码行。
    无文本层（扫描件）时返回空列表，由调用方报错提示 OCR。
    """
    from collections import Counter

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_blocks, line_freq, x0_freq = [], Counter(), Counter()
    for page in doc:
        page_dict = page.get_text("dict")
        blocks, seen_norms = [], set()
        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            lines = []
            for line in block.get("lines", []):
                text = "".join(span.get("text", "") for span in line.get("spans", []))
                if text.strip():
                    lines.append((text, line["bbox"][0]))
                    seen_norms.add(re.sub(r"\d+", "#", text.strip()))
                    x0_freq[round(line["bbox"][0])] += 1
            if lines:
                blocks.append(lines)
        pages_blocks.append(blocks)
        for norm in seen_norms:
            line_freq[norm] += 1
    doc.close()

    if not x0_freq:
        return []

    # 跨页高频重复的短行视为页眉/页脚样板文本
    boilerplate = {t for t, n in line_freq.items()
                   if n >= max(5, 0.2 * len(pages_blocks)) and len(t) < 80}
    body_x0 = x0_freq.most_common(1)[0][0]  # 正文左缘 = 出现最多的 x0

    paragraphs = []
    for blocks in pages_blocks:
        for lines in blocks:
            lines = [(t, x) for t, x in lines
                     if re.sub(r"\d+", "#", t.strip()) not in boilerplate
                     and not re.fullmatch(r"\d{1,4}", t.strip())
                     and not _ORNAMENT_RE.match(t.strip())]
            if not lines:
                continue
            # 首行缩进 -> 新段落
            groups, current = [], [lines[0][0]]
            for text, x0 in lines[1:]:
                if x0 >= body_x0 + 1.5:
                    groups.append(current)
                    current = [text]
                else:
                    current.append(text)
            groups.append(current)
            for group in groups:
                text = group[0].strip()
                for ln in group[1:]:
                    ln = ln.strip()
                    if text.endswith("-") and ln[:1].islower():
                        text = text[:-1] + ln  # 连字符换行修复
                    else:
                        text = text + " " + ln
                text = re.sub(r"\s+", " ", text).strip()
                if not text:
                    continue
                # 跨页/跨块延续：上一段未完结（且不是标题级短行）-> 合并
                if paragraphs and paragraphs[-1][-1] not in _SENTENCE_TERMINAL \
                        and len(paragraphs[-1]) > 40:
                    paragraphs[-1] = paragraphs[-1] + " " + text
                else:
                    paragraphs.append(text)

    # 兜底：小写开头的碎句并入上一段
    merged = []
    for para in paragraphs:
        if merged and para[:1].islower():
            merged[-1] = merged[-1] + " " + para
        else:
            merged.append(para)
    return [p for p in merged if not _ORNAMENT_RE.match(p)]


def call_llm(provider, api_key, model, system_prompt, user_prompt, temperature=0.1):
    """底层大模型统一路由（超时 150 秒，模型可配置）。"""
    if provider == "DeepSeek":
        client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com", timeout=150.0)
        res = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_prompt},
                      {"role": "user", "content": user_prompt}],
            temperature=temperature,
        )
        return (res.choices[0].message.content or "").strip()
    if provider == "OpenAI":
        client = OpenAI(api_key=api_key, timeout=150.0)
        res = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_prompt},
                      {"role": "user", "content": user_prompt}],
            temperature=temperature,
        )
        return (res.choices[0].message.content or "").strip()
    if provider == "Gemini":
        try:
            client = genai.Client(api_key=api_key,
                                  http_options=genai.types.HttpOptions(timeout=150_000))
        except (AttributeError, TypeError):
            client = genai.Client(api_key=api_key)
        res = client.models.generate_content(
            model=model,
            contents=user_prompt,
            system_instruction=system_prompt,
            config=genai.types.GenerateContentConfig(temperature=temperature),
        )
        return (res.text or "").strip()
    return ""


def parse_termbase(file_stream):
    """解析用户上传的术语库 Excel，返回概念化术语条目列表。

    必选列：Source / Target；可选列：Behavior（translate/preserve）、
    Status（locked/provisional）、Preferred（首选译名）、Forbidden（禁止译名，
    可用 ; 或 , 分隔）、Scope、Note。解析失败抛出 ValueError（不再静默吞错）。
    """
    try:
        df = pd.read_excel(file_stream)
    except Exception as e:
        raise ValueError(f"无法读取 Excel 文件：{e}") from e
    df.columns = [str(c).strip() for c in df.columns]
    if "Source" not in df.columns or "Target" not in df.columns:
        raise ValueError("术语库缺少 Source / Target 列，请检查表头")
    df = df.dropna(subset=["Source", "Target"])
    entries = []
    for _, row in df.iterrows():
        entry = {"source": str(row["Source"]).strip(),
                 "target": str(row["Target"]).strip()}
        for col, key in (("Behavior", "behavior"), ("Status", "status"),
                         ("Preferred", "preferred"), ("Scope", "scope"),
                         ("Note", "note")):
            if col in df.columns and pd.notna(row[col]):
                entry[key] = str(row[col]).strip()
        if "Forbidden" in df.columns and pd.notna(row["Forbidden"]):
            entry["forbidden"] = [x.strip() for x in
                                  re.split(r"[;；,，]", str(row["Forbidden"])) if x.strip()]
        entries.append(entry)
    return entries


def extract_auto_terms(paragraphs, target_lang, provider, api_key, model):
    """自动抽取术语库（兼容旧接口：返回 {source: target}）。

    新实现（mti_tool.terminology.extract_auto_terms_v2）：分布式采样、
    全量 occurrences、candidate 状态与 model_knowledge 证据。
    """
    from mti_tool.terminology import extract_auto_terms_v2
    entries, _warnings = extract_auto_terms_v2(
        paragraphs, target_lang, provider, api_key, model)
    return {e["source"]: e["target"] for e in entries}


# ================= 概念化术语表（对齐 localize-anything 的 Glossary 模型）=================
def normalize_glossary(entries):
    """标准化术语条目（委托 mti_tool.models，兼容旧字段并新增 id/occurrences/evidence）。"""
    return _models.normalize_glossary(entries)


def glossary_block(glossary):
    """把术语表渲染成注入翻译/审校 prompt 的文本块。"""
    locked_translate = [e for e in glossary if e["behavior"] == "translate" and e["status"] == "locked"]
    preserve = [e for e in glossary if e["behavior"] == "preserve"]
    provisional = [e for e in glossary if e["behavior"] == "translate" and e["status"] != "locked"]
    lines = []
    if locked_translate:
        lines.append("【锁定术语（必须使用首选译名，不得使用禁止译名）】：")
        for e in locked_translate:
            seg = f"- {e['source']} -> {e['preferred']}"
            if e["forbidden"]:
                seg += f"（禁止：{'、'.join(e['forbidden'])}）"
            lines.append(seg)
    if preserve:
        lines.append("【必须保留原文的术语/名称】：" + "、".join(e["source"] for e in preserve))
    if provisional:
        lines.append("【建议术语（仅供参考，请优先采用）】：")
        for e in provisional:
            lines.append(f"- {e['source']} -> {e['target']}")
    return "\n".join(lines)


def glossary_to_terms(glossary):
    """翻译行为术语 -> 扁平 dict（供报告生成等场景使用）。"""
    return {e["source"]: (e["preferred"] or e["target"])
            for e in glossary if e["behavior"] == "translate" and e["target"]}


def check_glossary_compliance(src, tgt, glossary):
    """锁定术语的确定性合规检查（委托 mti_tool.terminology：entry_id/segment_id 级）。"""
    from mti_tool.terminology import check_glossary_compliance as _qa
    return _qa(src, tgt, glossary)


# ================= 确定性检查（对齐 localize-anything 的机械检查）=================
PRESERVE_RE = re.compile(
    r'(?P<placeholder>%[sd]|%1\$[sd]|\{[A-Za-z_][A-Za-z0-9_]*\}|\{\{[A-Za-z_][A-Za-z0-9_]*\}\})'
    r'|(?P<url>https?://\S+|www\.\S+)'
    r'|(?P<email>[\w.+-]+@[\w-]+(?:\.[\w-]+)+)'
    r'|(?P<doi>10\.\d{4,9}/[^\s]+)'
    r'|(?P<citation>\[\d+(?:[-,]\s*\d+)*\])',
    re.IGNORECASE,
)

PRESERVE_SEVERITY = {
    "placeholder": "blocking",   # 占位符损坏 = 结构破坏，绝不可自动放行
    "url": "actionable",
    "email": "actionable",
    "doi": "actionable",
    "citation": "actionable",
}


def extract_preserved_tokens(text):
    """提取源文本中必须原样保留的 token（占位符/URL/邮箱/DOI/引用标注）。"""
    return {m.group(0): m.lastgroup for m in PRESERVE_RE.finditer(text or "")}


def find_residuals(src, tgt, target_lang):
    """检测目标语言中残留的源语言片段。

    返回 [(片段, severity)]：连续 ≥2 个源语单词/较长汉字串 -> actionable；
    单个词（可能是专有名词）-> informational。启发式，不替代审校。
    """
    tgt_clean = PRESERVE_RE.sub(" ", tgt or "")
    if target_lang == "English":
        source_runs = set(re.findall(r'[\u4e00-\u9fff]{2,}', src or ""))
        return [(c, "actionable" if len(c) >= 4 else "informational")
                for c in re.findall(r'[\u4e00-\u9fff]{2,}', tgt_clean)
                if any(c in run for run in source_runs)]
    src_words = set(w.lower() for w in re.findall(r'[A-Za-z]{5,}', src or ""))
    allowed = {"mti"}  # 产品名等明确保留词白名单（审校负责语义判断）
    words = re.findall(r'[A-Za-z]{5,}', tgt_clean)
    hits = [w for w in words if w.lower() in src_words and w.lower() not in allowed]
    result, run = [], []
    for w in words:
        if w in hits:
            run.append(w)
        else:
            if run:
                result.append((" ".join(run),
                               "actionable" if len(run) >= 2 else "informational"))
                run = []
    if run:
        result.append((" ".join(run), "actionable" if len(run) >= 2 else "informational"))
    return result


def _count_sentences(text):
    """粗粒度句数统计：按终结符切分（引号/括号闭合归并到前一句）。"""
    text = _ABBREV_RE.sub(" ", text)
    parts = re.split(r"[.!?…。！？]+[”\"'’)\]]*", text)
    return sum(1 for p in parts if p.strip())


def is_incomplete_translation(src, tgt):
    """疑似漏译/截断判定（双重规则，实测调优）：
    1. 字符级：长原文（≥120 字符）配极短译文（<15%）——只拦灾难性截断，
       英译中正常比例可低至 0.2-0.3，不能用高阈值；
    2. 句子级：原文 ≥2 句而译文不足一半句数，且字符占比 <35%——
       截断译文必然句数对不上，完整译文即使语言再凝练也很少掉一半句。
    """
    tgt = (tgt or "").strip()
    if not tgt:
        return True
    if len(src) >= 120 and len(tgt) < 0.15 * len(src):
        return True
    src_sents = _count_sentences(src)
    if src_sents >= 2:
        tgt_sents = _count_sentences(tgt)
        if tgt_sents < src_sents * 0.5 and len(tgt) < 0.35 * len(src):
            return True
    return False


def check_translation_batch(sources, targets, glossary, target_lang):
    """确定性检查一批译文：空译、保留项丢失、源语残留、锁定术语合规。"""
    findings = []
    for i, (src, tgt) in enumerate(zip(sources, targets)):
        if not tgt.strip():
            findings.append({"segment_index": i, "type": "check", "severity": "blocking",
                             "reason": "译文为空"})
            continue
        # 完整性检查：拦截截断译文。
        # 实测根因：审校/修复环节的整段替换把长段译文换成了一句修正。
        if is_incomplete_translation(src, tgt):
            findings.append({"segment_index": i, "type": "check", "severity": "blocking",
                             "reason": f"疑似漏译/截断：原文 {len(src)} 字符"
                                       f"/{_count_sentences(src)} 句，译文仅 {len(tgt.strip())} 字符"
                                       f"/{_count_sentences(tgt)} 句"})
        for token, kind in extract_preserved_tokens(src).items():
            if token not in tgt:
                findings.append({"segment_index": i, "type": "check",
                                 "severity": PRESERVE_SEVERITY.get(kind, "actionable"),
                                 "reason": f"保留项 {kind}「{token}」在译文中丢失"})
        for residual, sev in find_residuals(src, tgt, target_lang):
            findings.append({"segment_index": i, "type": "check", "severity": sev,
                             "reason": f"疑似残留源语片段「{residual}」"})
        findings.extend(check_glossary_compliance(src, tgt, glossary))
        for f in findings:
            if "segment_index" not in f:
                f["segment_index"] = i
            if "segment_id" not in f or f.get("segment_id") is None:
                f["segment_id"] = i
    return findings


# ================= 语义批次（对齐 localize-anything 的上下文批次）=================
BATCH_SIZE = 4
MAX_BATCH_CHARS = 1600


def make_batches(paragraphs, batch_size=BATCH_SIZE, max_chars=MAX_BATCH_CHARS):
    """把段落聚成语义批次：≤batch_size 段且累计 ≤max_chars 字符。"""
    batches, cur, n = [], [], 0
    for p in paragraphs:
        if cur and (len(cur) >= batch_size or n + len(p) > max_chars):
            batches.append(cur)
            cur, n = [], 0
        cur.append(p)
        n += len(p)
    if cur:
        batches.append(cur)
    return batches


def _batch_section_profile(document_profile, offset, batch_len):
    """按全局段区间匹配 section profile（用于相关术语的 section:<id> scope）。"""
    if not document_profile:
        return None
    for sec in document_profile.get("sections") or []:
        start, end = sec.get("start_segment"), sec.get("end_segment")
        if start is None or end is None:
            continue
        if start <= offset and offset + batch_len - 1 <= end:
            return sec
    return None


# ================= 翻译记忆（对齐 localize-anything 的 TM：仅收录审校通过段落）=================
def tm_path():
    return OUTPUT_DIR / "translation_memory.json"


def _tm_eligible(source, target):
    """翻译记忆资格：源文必须有字母/数字（纯符号装饰行不入库），译文非空。"""
    return bool(re.search(r"[A-Za-z0-9\u4e00-\u9fff]", source or "")) \
        and bool((target or "").strip())


def load_tm():
    """加载翻译记忆并自清洗：非法条目（无字母源文/空译文/未过审校）直接丢弃。

    翻译记忆是错误放大器（一次错译会复制到全书），因此加载即消毒，
    防止旧版本或异常写入留下的污染条目继续命中。
    """
    p = tm_path()
    if p.is_file():
        try:
            raw = json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            return {}
        return {k: v for k, v in raw.items()
                if isinstance(v, dict) and v.get("reviewed")
                and _tm_eligible(k, v.get("target"))}
    return {}


def save_tm(tm):
    p = tm_path()
    p.parent.mkdir(parents=True, exist_ok=True)
    tmp = p.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(tm, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(p)


# ================= 翻译 / 修复 / 审校（对齐 localize-anything 的三通道）=================
def _translator_system(glossary_text, style_rules, target_lang):
    return (f"你是一位学术翻译专家，请将用户提供的段落翻译成{target_lang}。\n"
            f"规则：只翻译可翻译的正文；作者姓名、机构名、品牌名、URL、邮箱、DOI、"
            f"引用标注（如 [12]）等保留原文；译文须与原文一一对应并保持顺序。\n"
            f"{glossary_text}\n"
            f"{style_rules}\n"
            "请严格输出合法的 JSON 字符串数组，不要包含任何解释文字。")


def translate_batch(segments, ctx_prev, ctx_next, glossary_text, style_rules, target_lang,
                    provider, api_key, model):
    """翻译一个语义批次，返回与 segments 等长的译文列表；失败抛出 RuntimeError。"""
    numbered = "\n".join(f"{i + 1}. {s}" for i, s in enumerate(segments))
    context = ""
    if ctx_prev:
        context += "【前文上下文】：\n" + "\n".join(f"- {s}" for s in ctx_prev) + "\n\n"
    if ctx_next:
        context += "【后文上下文】：\n" + "\n".join(f"- {s}" for s in ctx_next) + "\n\n"
    sys_prompt = _translator_system(glossary_text, style_rules, target_lang)
    user_prompt = f"{context}待翻译段落（按序号返回等长译文数组）：\n{numbered}"
    last_err, last_res = None, None
    for _attempt in range(3):
        try:
            res = call_llm(provider, api_key, model, sys_prompt, user_prompt, temperature=0.3)
            last_res = res
            arr = parse_translation_array(res, len(segments))
            if arr is None:
                raise ValueError(f"译文数量不匹配：期望 {len(segments)}，"
                                 f"响应预览：{(res or '')[:160]!r}")
            # 空项不直接判死：交给确定性检查与自动修复环节兜底
            return [item.strip() if isinstance(item, str) else "" for item in arr]
        except Exception as e:
            last_err = e
            if is_rate_limited(e):
                time.sleep(15)
            else:
                break
    if len(segments) == 1 and last_res and str(last_res).strip():
        # 单段兜底：接受模型原始输出，交由确定性检查/修复把关
        return [re.sub(r'^\s*\d+[.)、]\s*', '', str(last_res)).strip()]
    raise RuntimeError(f"批次翻译失败：{last_err or '模型返回格式异常或数量不匹配'}")


def repair_batch(sources, targets, findings, glossary_text, style_rules, target_lang,
                 provider, api_key, model):
    """根据确定性检查发现的问题自动修复一批译文；返回与 sources 等长的译文列表。"""
    numbered = "\n".join(
        f"{i + 1}. 原文：{s}\n   译文：{t}" for i, (s, t) in enumerate(zip(sources, targets)))
    issues = "\n".join(f"- 段落 {f['segment_index'] + 1}: [{f['severity']}] {f['reason']}"
                       for f in findings)
    sys_prompt = _translator_system(glossary_text, style_rules, target_lang)
    user_prompt = ("以下译文未通过检查，请仅修正有问题的段落，其余段落保持原样，"
                   f"返回与段落数相同的 JSON 字符串数组：\n\n{numbered}\n\n问题清单：\n{issues}")
    for _attempt in range(3):
        try:
            res = call_llm(provider, api_key, model, sys_prompt, user_prompt, temperature=0.2)
            arr = parse_json_array(res)
            if arr is None or len(arr) != len(sources):
                raise ValueError("修复结果数量不匹配")
            return [clean_xml_chars(str(x)).strip() if isinstance(x, str) else "" for x in arr]
        except Exception as e:
            if is_rate_limited(e):
                time.sleep(15)
            else:
                break
    raise RuntimeError("自动修复失败：模型返回格式异常")


def review_translation_batch(sources, targets, glossary_text, style_rules, target_lang,
                             provider, api_key, model):
    """独立审校一个批次（与翻译分离的 prompt/上下文），返回 (findings, failed)。"""
    numbered = "\n".join(
        f"{i + 1}. 原文：{s}\n   译文：{t}" for i, (s, t) in enumerate(zip(sources, targets)))
    sys_prompt = (f"你是一位独立的翻译审校专家，负责审查机器译文。请检查：语义准确性、术语一致性、"
                  f"漏译/增译、目标语言自然度与风格。只报告真实存在的问题，"
                  f"不要为低风险或主观偏好制造 finding。\n"
                  f"severity 只允许以下三种：blocking（结构/占位符/语义严重错误）、"
                  f"actionable（应修正的问题）、informational（建议）。\n"
                  "如果整批译文没有问题，请严格返回空数组 []，不要输出任何 informational 备注。\n"
                  f"{glossary_text}\n"
                  f"{style_rules}\n"
                  '请严格输出 JSON 数组，每项格式：{"segment_index": 0, "severity": "actionable", '
                  '"reason": "问题说明", "suggested_target": "可选：修正后的译文"}')
    user_prompt = f"待审校段落（目标语言：{target_lang}）：\n{numbered}"
    for _attempt in range(3):
        try:
            res = call_llm(provider, api_key, model, sys_prompt, user_prompt, temperature=0.2)
            arr = parse_json_array(res)
            if arr is None:
                return [], True
            return arr, False
        except Exception as e:
            if is_rate_limited(e):
                time.sleep(15)
            else:
                break
    return [], True


# ================= 自动标注（三色学习重点）=================
# 红=生僻词/难词；黄=专业名词（特殊译法）；青绿=翻译难点句（特别译法）。
ANNOT_BATCH_SIZE = 10
ANNOT_MAX_PER_SEG = {"rare": 3, "domain": 3, "hard": 2}

# 常用英语词表（en_50k 字幕语料前 14000 词），用于把 LLM 滥标的"生僻词"挡回去
_DATA_DIR = Path(__file__).resolve().parent / "data"
_COMMON_WORDS = None


def _common_words():
    global _COMMON_WORDS
    if _COMMON_WORDS is None:
        try:
            _COMMON_WORDS = set(
                _DATA_DIR.joinpath("en_common.txt").read_text(encoding="utf-8").splitlines())
        except OSError:
            _COMMON_WORDS = set()
    return _COMMON_WORDS


_INFLECTION_SUFFIXES = (("ily", "y"), ("ness", ""), ("ment", ""), ("tion", ""),
                        ("sion", ""), ("ing", ""), ("ed", ""), ("er", ""),
                        ("est", ""), ("es", ""), ("ly", ""), ("s", ""))


def _base_form(word):
    """词形还原（一次）：grinning->grin、speedily->speedy、boxes->box。"""
    w = word
    for suffix, repl in _INFLECTION_SUFFIXES:
        if w.endswith(suffix) and len(w) > len(suffix) + 2:
            w = w[:-len(suffix)] + repl
            break
    if len(w) > 3 and w[-1] == w[-2] and w[-1] not in "eio":
        w = w[:-1]  # grinning -> grinn -> grin
    return w


def _is_common_word(token):
    """token 或其词形还原是否在常用词表内。"""
    common = _common_words()
    if not common:
        return False
    w = token.lower().strip('"\'(),;:!?')
    if not w:
        return True
    if w in common:
        return True
    base = _base_form(w)
    return base != w and base in common


_KINSHIP_TITLE_RE = re.compile(
    r"^(?:grandma|grandpa|grandmother|grandfather|aunt|uncle|mr|mrs|ms|dr|sir|lady|lord|"
    r"mother|father|mom|dad|brother|sister|captain|colonel|major|general|rabbi|"
    r"professor|doctor)\b", re.IGNORECASE)


def _rare_ok(span_text, token_freq):
    """生僻词门槛：单 token、非常用词（含词形还原）、全书出现次数少。"""
    span = (span_text or "").strip()
    if not span or " " in span or "\u00a0" in span:
        return False  # 只接受单词（可带连字符），短语/名句一律不要
    w = span.lower().strip('"\'(),;:!?')
    if len(w) < 4:
        return False
    if _is_common_word(w):
        return False
    if token_freq.get(w, 0) >= 8:
        return False  # 全书反复出现的词不算生僻
    return True


def _domain_ok(span_text):
    """专业名词门槛：称谓+人名、全常用词短语不算专业名词。"""
    span = (span_text or "").strip()
    if not span:
        return False
    if _KINSHIP_TITLE_RE.match(span):
        return False  # Grandma Sarah / Mr. Smith 之类
    tokens = re.findall(r"[A-Za-z]+(?:['’\-][A-Za-z]+)*", span)
    if not tokens:
        return False
    if all(_is_common_word(t) for t in tokens):
        return False  # Translation from Hebrew 之类全是常用词
    return True


def _normalize_annotations(annotations):
    """标注字典键归一化：JSON 落盘后键变字符串，统一回 int；越界/非列表值丢弃。

    这是本项目反复踩过的坑（标注渲染、过滤、续跑三处各自处理过一次），
    统一入口避免再次各修各的。
    """
    out = {}
    for k, v in (annotations or {}).items():
        try:
            gi = int(k)
        except (TypeError, ValueError):
            continue
        if isinstance(v, list):
            out[gi] = v
    return out


def _clean_annotations(annotations, pairs):
    """过滤 + 去重 + 数量上限；术语表覆盖的 domain（note 以"术语："开头）不参与过滤。"""
    token_freq = {}
    for pr in pairs:
        for tok in re.findall(r"[A-Za-z]+", pr["source"].lower()):
            token_freq[tok] = token_freq.get(tok, 0) + 1
    cleaned = {}
    for gi, items in _normalize_annotations(annotations).items():
        if not (0 <= gi < len(pairs)):
            continue  # 段已被删除等场景：越界键丢弃
        src_text = pairs[gi]["source"]
        seen, kept, counts = set(), [], {"rare": 0, "domain": 0, "hard": 0}
        for it in items:
            atype = it["type"]
            s_span = it.get("src_span")
            span_text = src_text[s_span[0]:s_span[1]] if s_span else ""
            overlay = atype == "domain" and str(it.get("note", "")).startswith("术语：")
            if atype == "rare" and not _rare_ok(span_text, token_freq):
                continue
            if atype == "domain" and not overlay and not _domain_ok(span_text):
                continue
            key = (atype, tuple(s_span) if s_span else None,
                   tuple(it["tgt_span"]) if it.get("tgt_span") else None)
            if key in seen:
                continue
            seen.add(key)
            if counts[atype] >= ANNOT_MAX_PER_SEG[atype]:
                continue
            counts[atype] += 1
            kept.append(it)
        if kept:
            cleaned[gi] = kept
    return cleaned


def _annotator_system(target_lang):
    return (f"你是一位翻译教学专家。请从下列{target_lang}双语对照中标注三类学习重点：\n"
            "1. rare：真正的生僻词/难词——英语母语者也未必认识的低频书面词，"
            "如 chicory、muezzin、cacophony。只标注单个单词（最多一个带连字符的复合词）。\n"
            "   严禁标注日常常用词（如 production、grin、rooster、speedily、elementary），"
            "严禁标注短语、引用句或名句（如 'Elementary, my dear Watson'）。\n"
            "2. domain：专业领域名词，译文采用了专门/约定俗成的译法（如术语表、行话、专名译法）；\n"
            "   严禁标注亲属称谓（如 Grandma Sarah）、全常用词短语（如 Translation from Hebrew）、"
            "普通日常表达。\n"
            "3. hard：翻译难度高的句子，译文使用了特别翻译技巧（语序调整、词性转换、拆合句、"
            "文化负载词处理、比喻/双关处理等）；普通直译句不要标注。\n"
            "规则：\n"
            "- 每个段落最多 2 个 rare、2 个 domain、1 个 hard；标注真正有价值的，宁缺毋滥；\n"
            "- src 必须是原文中的原文字符串，tgt 必须是对应译文中的字符串（hard 可以是整句/整段）；\n"
            "- note 用一句话说明标注理由或所用译法；\n"
            '严格输出 JSON 数组，每项格式：{"seg": 1, "type": "rare", "src": "...", '
            '"tgt": "...", "note": "..."}。seg 为段落序号（从 1 开始）。不要输出任何解释文字。')


def annotate_batch(pairs_slice, target_lang, provider, api_key, model):
    """标注一个批次，返回 [{seg, type, src, tgt, note}]（seg 为 0 基）；解析失败返回 []。"""
    numbered = "\n\n".join(
        f"--- 段落 {i + 1} ---\n原文：{p['source']}\n译文：{p['target']}"
        for i, p in enumerate(pairs_slice))
    sys_prompt = _annotator_system(target_lang)
    user_prompt = f"待标注双语段落：\n{numbered}"
    for _attempt in range(3):
        try:
            res = call_llm(provider, api_key, model, sys_prompt, user_prompt, temperature=0.2)
            arr = parse_json_array(res)
            if arr is None:
                return []
            out = []
            for item in arr:
                if not isinstance(item, dict):
                    continue
                seg = item.get("seg")
                if not isinstance(seg, int) or not (1 <= seg <= len(pairs_slice)):
                    continue
                atype = item.get("type")
                if atype not in ANNOTATION_COLORS:
                    continue
                src = str(item.get("src") or "").strip()
                tgt = str(item.get("tgt") or "").strip()
                note = str(item.get("note") or "").strip()
                if not src:
                    continue
                out.append({"seg": seg - 1, "type": atype, "src": src,
                            "tgt": tgt, "note": note})
            return out
        except Exception as e:
            if is_rate_limited(e):
                time.sleep(10)  # 限流退避后重试，避免整批标注静默丢失
                continue
            return []
    return []


def _compose_spans(spans, text_len):
    """按边界切分 + 优先级（rare>domain>hard）覆盖，返回互不重叠的 (start,end,type) 列表。

    难点句常覆盖整段，词级标注嵌在其中：切到所有起点/终点后，每段取覆盖它的最高优先级。
    """
    priority = {"rare": 0, "domain": 1, "hard": 2}
    clipped = []
    for s, e, t in spans:
        s = max(0, min(int(s), text_len))
        e = max(s, min(int(e), text_len))
        if s < e:
            clipped.append((s, e, t))
    if not clipped:
        return []
    bounds = sorted({0, text_len} | {x for s, e, _ in clipped for x in (s, e)})
    composed = []
    for a, b in zip(bounds, bounds[1:]):
        if a >= b:
            continue
        mid = (a + b) / 2
        covering = [(priority[t], t) for s, e, t in clipped if s <= mid < e]
        if covering:
            composed.append((a, b, min(covering)[1]))
    return composed


def annotate_stage(state, job_id, glossary, provider, api_key, model, target_lang,
                   on_caption=None):
    """三色自动标注：LLM 识别 + 术语表确定性覆盖（专业名词=黄色必标）。"""
    if state.get("annotations_done"):
        return state
    pairs = state["pairs"]
    annotations = _normalize_annotations(state.get("annotations"))
    batches = make_batches([p["source"] for p in pairs], batch_size=ANNOT_BATCH_SIZE,
                           max_chars=2600)
    # 断点按“已标注段数”记录（而非批次号），避免批大小调整后续跑错位
    done_offset = state.get("annotations_done_offset") or 0
    start_bi = next((bi for bi, b in enumerate(batches)
                     if sum(len(x) for x in batches[:bi]) >= done_offset), len(batches))
    if done_offset:
        if on_caption:
            on_caption(f"↩️ 从第 {done_offset + 1} 段（批次 {start_bi + 1}/{len(batches)}）继续标注...")

    # 1) LLM 标注
    failed_batches = 0
    for bi in range(start_bi, len(batches)):
        batch_srcs = batches[bi]
        offset = sum(len(b) for b in batches[:bi])
        slice_pairs = pairs[offset:offset + len(batch_srcs)]
        if on_caption and bi % 10 == 0:
            on_caption(f"🎨 自动标注第 {offset + 1}-{offset + len(slice_pairs)} 段"
                       f"（共 {len(pairs)} 段，批次 {bi + 1}/{len(batches)}）...")
        items = annotate_batch(slice_pairs, target_lang, provider, api_key, model)
        if not items:
            failed_batches += 1
        for item in items:
            gi = offset + item["seg"]
            src, tgt, atype, note = item["src"], item["tgt"], item["type"], item["note"]
            src_span = _find_span(pairs[gi]["source"], src)
            tgt_span = _find_span(pairs[gi]["target"], tgt) if tgt else None
            if atype == "hard":
                # 难句兜底：找不到精确片段时标整段
                if src_span is None:
                    src_span = (0, len(pairs[gi]["source"]))
                if tgt_span is None:
                    tgt_span = (0, len(pairs[gi]["target"])) if pairs[gi]["target"] else None
            elif src_span is None:
                continue  # 词级标注必须在原文中定位
            annotations.setdefault(gi, []).append(
                {"type": atype, "src_span": list(src_span) if src_span else None,
                 "tgt_span": list(tgt_span) if tgt_span else None, "note": note})
        # 每批落盘：断点粒度 = 一个批次
        state["annotations"] = annotations
        state["annotations_done_offset"] = offset + len(slice_pairs)
        save_job_state(job_id, state)

    # 2) 术语表确定性覆盖：专业名词（特殊译法）-> 黄色
    for gi, pr in enumerate(pairs):
        for entry in glossary:
            term = (entry.get("source") or "").strip()
            if len(term) < 2 or term not in pr["source"]:
                continue
            span = _find_span(pr["source"], term)
            tgt_span = None
            tgt_term = (entry.get("target") or "").strip()
            if tgt_term:
                tgt_span = _find_span(pr["target"], tgt_term)
            annotations.setdefault(gi, []).append(
                {"type": "domain", "src_span": list(span) if span else None,
                 "tgt_span": list(tgt_span) if tgt_span else None,
                 "note": f"术语：{term} -> {tgt_term or '保留原文'}"})

    # 3) 确定性过滤（常用词/称谓/全常用词短语）+ 去重 + 数量上限
    cleaned = _clean_annotations(annotations, pairs)
    state["annotations"] = cleaned
    state["annotations_done"] = True
    state["annotations_failed_batches"] = failed_batches
    save_job_state(job_id, state)
    if on_caption:
        total = sum(len(v) for v in cleaned.values())
        on_caption(f"✅ 自动标注完成：{total} 处（失败批次 {failed_batches}/{len(batches)}）")
    return state


def translate_stage(state, job_id, glossary, provider, api_key, model, target_lang,
                    style_rules, enable_review, document_profile=None,
                    on_status=None, on_caption=None):
    """阶段二：语义批次翻译 + 确定性检查/修复 + 独立审校 + 翻译记忆。

    对齐 localize-anything 经验：
    - 语义批次：≤4 段一组，携带前后文，保留每批落盘的断点粒度；
    - 概念化术语表：锁定术语强制首选译名/禁止译名，保留项强制原样；
    - 确定性检查：占位符/URL/引用等保留项、残留原文、锁定术语合规，问题自动修复一轮；
    - 独立审校：actionable 建议经确定性复验后应用，blocking 记录给用户确认；
    - 翻译记忆：仅审校通过的段落入库，精确命中直接复用。
    """
    tm = load_tm()
    paras = state["paras"]
    pairs = state["pairs"]
    batches = make_batches(paras)

    # 断点：从第一个未完成批次继续；若中间批次不完整则截断重译
    cum_end, start_batch = 0, 0
    for bi, b in enumerate(batches):
        prev_end = cum_end
        cum_end += len(b)
        if cum_end > len(pairs):
            if len(pairs) > prev_end:
                del pairs[prev_end:]
            start_batch = bi
            break
    else:
        start_batch = len(batches)

    stats = state.setdefault("review_stats", {
        "reviewed_segments": 0, "batches_reviewed": 0,
        "blocking": 0, "actionable": 0, "informational": 0, "review_failed": 0,
    })
    findings_all = state.setdefault("findings", [])
    from mti_tool.terminology import (
        detect_glossary_conflicts as _detect_conflicts,
        glossary_block as _glossary_block,
        select_glossary_for_segments as _select_glossary,
    )

    for bi in range(start_batch, len(batches)):
        batch = batches[bi]
        offset = sum(len(b) for b in batches[:bi])
        if on_status:
            on_status(f"【阶段二】双语翻译与术语严格注入...（批次 {bi + 1}/{len(batches)}）")
        if on_caption:
            on_caption(f"🌍 正在翻译第 {offset + 1}-{offset + len(batch)} 段（共 {len(paras)} 段）...")

        ctx_prev = paras[max(0, offset - 2):offset]
        ctx_next = paras[min(len(paras), offset + len(batch)):min(len(paras), offset + len(batch) + 2)]

        # 1) 翻译记忆精确命中直接复用
        batch_pairs = [None] * len(batch)
        to_translate = []  # (index, clean_source)
        for i, para in enumerate(batch):
            clean_src = para.replace('\n', ' ')
            hit = tm.get(clean_src)
            if hit and hit.get("reviewed") and hit.get("target"):
                batch_pairs[i] = {"source": clean_src, "target": hit["target"],
                                  "reviewed": True, "from_tm": True}
                state["tm_used_count"] = state.get("tm_used_count", 0) + 1
            elif not re.search(r"[A-Za-z0-9\u4e00-\u9fff]", clean_src):
                # 纯符号段落（章节分隔装饰等）：不是正文，原样保留，不调模型
                batch_pairs[i] = {"source": clean_src, "target": clean_src,
                                  "reviewed": True, "from_tm": False}
            else:
                to_translate.append((i, clean_src))

        # 相关术语选择：只注入本批实际出现的 locked translate / preserve 条目，
        # provisional 仅作受限建议；记录实际注入的 entry IDs（审计）。
        texts = [t for _, t in to_translate]
        section_profile = _batch_section_profile(document_profile, offset, len(batch))
        selected, injected_ids = _select_glossary(
            texts, glossary, document_profile, section_profile)
        glossary_text = _glossary_block(selected)
        state.setdefault("glossary_injection_log", []).append({
            "batch": bi,
            "offset": offset,
            "entry_ids": injected_ids,
            "glossary_version": (state.get("glossary_frozen") or {}).get("version"),
            "glossary_hash": (state.get("glossary_frozen") or {}).get("glossary_hash"),
        })

        # 2) 未命中段落批次翻译
        if to_translate:
            try:
                targets = translate_batch(texts, ctx_prev, ctx_next, glossary_text, style_rules,
                                          target_lang, provider, api_key, model)
            except RuntimeError:
                # 批次解析失败时降级为逐段翻译，保证进度不中断
                if len(texts) == 1:
                    raise
                if on_caption:
                    on_caption("⚠️ 批次翻译返回格式异常，降级为逐段翻译...")
                targets = []
                for t in texts:
                    targets.append(translate_batch([t], ctx_prev, ctx_next, glossary_text,
                                                   style_rules, target_lang, provider,
                                                   api_key, model)[0])
            for (i, src), tgt in zip(to_translate, targets):
                batch_pairs[i] = {"source": src, "target": clean_xml_chars(tgt).replace('\n', ' '),
                                  "reviewed": False, "from_tm": False}

        batch_sources = [p["source"] for p in batch_pairs]
        batch_targets = [p["target"] for p in batch_pairs]
        for p in batch_pairs:
            p["glossary_entry_ids"] = list(injected_ids)
        findings = check_translation_batch(batch_sources, batch_targets, glossary, target_lang)

        # 3) 确定性问题自动修复（一轮）
        fixable = [f for f in findings if f["severity"] in ("blocking", "actionable")]
        if fixable and len(fixable) <= 8:
            if on_caption:
                on_caption(f"🔧 发现 {len(fixable)} 个确定性问题，正在自动修复...")
            try:
                repaired = repair_batch(batch_sources, batch_targets, fixable, glossary_text,
                                        style_rules, target_lang, provider, api_key, model)
                if repaired and len(repaired) == len(batch_pairs):
                    for j, p in enumerate(batch_pairs):
                        if not p["from_tm"] and repaired[j] and repaired[j].strip():
                            candidate = clean_xml_chars(repaired[j]).replace('\n', ' ')
                            # 修复结果本身截断时，不接受更差的译文
                            if is_incomplete_translation(batch_sources[j], candidate) \
                                    and not is_incomplete_translation(batch_sources[j], p["target"]):
                                continue
                            p["target"] = candidate
                batch_targets = [p["target"] for p in batch_pairs]
                findings = check_translation_batch(batch_sources, batch_targets, glossary, target_lang)
            except Exception:
                pass  # 修复失败则保留原译文与 finding

        # 4) 独立审校：actionable 建议复验后应用；blocking 记录给用户
        if enable_review:
            stats["batches_reviewed"] += 1
            rfindings, failed = review_translation_batch(
                batch_sources, batch_targets, glossary_text, style_rules, target_lang,
                provider, api_key, model)
            if failed:
                stats["review_failed"] += 1
            for rf in rfindings:
                sev = rf.get("severity")
                if sev not in ("blocking", "actionable", "informational"):
                    continue
                idx = rf.get("segment_index")
                if not isinstance(idx, int) or not (0 <= idx < len(batch_pairs)):
                    continue
                record = {"segment_index": offset + idx, "severity": sev, "type": "review",
                          "reason": str(rf.get("reason") or "审校发现问题")}
                if sev == "actionable" and rf.get("suggested_target") \
                        and not batch_pairs[idx]["from_tm"]:
                    suggested = clean_xml_chars(rf["suggested_target"]).replace('\n', ' ').strip()
                    if suggested:
                        old_target = batch_pairs[idx]["target"]
                        batch_pairs[idx]["target"] = suggested
                        recheck = check_translation_batch(
                            [batch_sources[idx]], [suggested], glossary, target_lang)
                        if any(f["severity"] in ("blocking", "actionable") for f in recheck):
                            batch_pairs[idx]["target"] = old_target  # 复验不过则回滚
                            record["suggested_target"] = suggested
                            findings_all.append(record)
                        continue
                findings_all.append(record)

        # 审校可能修改过译文：对最终译文整体复验一次确定性检查
        findings = check_translation_batch(
            batch_sources, [p["target"] for p in batch_pairs], glossary, target_lang)

        # 批内冲突检测（跨段同术语多译法）——在 TM 入库前执行
        for cf in _detect_conflicts(batch_pairs, glossary):
            cf["segment_index"] = offset + cf["segment_id"]
            findings_all.append(cf)

        # 记录仍未解决的确定性问题
        for f in findings:
            if f["severity"] in ("blocking", "actionable", "informational"):
                findings_all.append({"segment_index": offset + f["segment_index"],
                                     "severity": f["severity"], "type": "check",
                                     "reason": f["reason"]})

        # 5) 审校通过的段落 -> 翻译记忆
        if enable_review:
            for j, p in enumerate(batch_pairs):
                seg_findings = [f for f in findings_all if f.get("segment_index") == offset + j
                                and f["severity"] in ("blocking", "actionable")]
                if not seg_findings and not p["from_tm"] \
                        and _tm_eligible(p["source"], p["target"]):
                    p["reviewed"] = True
                    tm[p["source"]] = {"target": p["target"], "reviewed": True}
                    stats["reviewed_segments"] += 1
            save_tm(tm)

        pairs.extend(batch_pairs)
        save_job_state(job_id, state)  # 每批落盘，断点粒度 = 一个批次

    # 全局冲突检测（跨批次），与批内结果去重
    batch_conflict_keys = {(f.get("type"), f.get("entry_id"),
                            f.get("segment_index"), True)
                           for f in findings_all if f.get("conflict")}
    for cf in _detect_conflicts(pairs, glossary):
        cf["segment_index"] = cf["segment_id"]
        key = (cf.get("type"), cf.get("entry_id"), cf.get("segment_index"), True)
        if key not in batch_conflict_keys:
            findings_all.append(cf)

    stats["blocking"] = sum(1 for f in findings_all if f["severity"] == "blocking")
    stats["actionable"] = sum(1 for f in findings_all if f["severity"] == "actionable")
    stats["informational"] = sum(1 for f in findings_all if f["severity"] == "informational")
    state["has_blocking"] = stats["blocking"] > 0
    return state


def findings_report_md(state):
    """把审查结果渲染成 Markdown 报告（下载/展示用）。"""
    stats = state.get("review_stats") or {}
    lines = [
        "# 翻译审查报告", "",
        "## 概览",
        f"- 已审校段落：{stats.get('reviewed_segments', 0)}",
        f"- 审校批次：{stats.get('batches_reviewed', 0)}",
        f"- 审校失败批次：{stats.get('review_failed', 0)}",
        f"- 翻译记忆复用：{state.get('tm_used_count', 0)} 段",
        f"- blocking：{stats.get('blocking', 0)}",
        f"- actionable：{stats.get('actionable', 0)}",
        f"- informational：{stats.get('informational', 0)}",
        "", "## 待处理问题",
    ]
    findings = state.get("findings") or []
    if not findings:
        lines.append("无。")
    else:
        for f in findings:
            line = f"- 第 {f.get('segment_index', -1) + 1} 段 [{f.get('severity')}] {f.get('reason')}"
            if f.get("suggested_target"):
                line += f"（建议译文：{f['suggested_target']}）"
            lines.append(line)
    return "\n".join(lines)


# ================= 文档/表格生成 =================
EN_FONT = "Times New Roman"
CN_FONT = "宋体"

# 自动标注三色：生僻词=红、专业名词（特殊译法）=黄、翻译难点句=青绿
ANNOTATION_COLORS = {"rare": "C00000", "domain": "BF8F00", "hard": "008080"}
ANNOTATION_LABELS = {"rare": "生僻词/难词", "domain": "专业名词（特殊译法）",
                     "hard": "翻译难点句（特别译法）"}


def _apply_doc_fonts(doc):
    """默认字体：西文 Times New Roman，中文宋体（Normal + 标题样式一并设置）。"""
    from docx.oxml.ns import qn
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Title"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = EN_FONT
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), EN_FONT)
        rfonts.set(qn("w:hAnsi"), EN_FONT)
        rfonts.set(qn("w:eastAsia"), CN_FONT)


def _apply_run_fonts(run):
    """单个 run 的字体（表格单元格里的 run 不受 Normal 样式继承影响时兜底）。"""
    from docx.oxml.ns import qn
    run.font.name = EN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), CN_FONT)


def dict_to_excel(term_dict):
    df = pd.DataFrame(list(term_dict.items()), columns=["Source", "Target"])
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    return output


def paragraphs_to_word(paragraphs):
    doc = Document()
    _apply_doc_fonts(doc)
    doc.add_heading('阶段一：清洗后原文提取', 0)
    for p in paragraphs:
        doc.add_paragraph(p)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _find_span(text, needle):
    """宽容定位子串：统一引号/破折号/省略号、折叠空白后查找，返回 (start, end) 或 None。"""
    if not needle or not text:
        return None
    if needle in text:
        pos = text.find(needle)
        return pos, pos + len(needle)
    mapping = []
    norm_chars = []
    for orig_idx, ch in enumerate(text):
        ch2 = ch.replace("’", "'").replace("‘", "'").replace("“", '"').replace("”", '"')
        ch2 = ch2.replace("–", "-").replace("—", "-").replace("…", "...")
        if ch2.isspace():
            if norm_chars and norm_chars[-1] != " ":
                norm_chars.append(" ")
                mapping.append(None)
            continue
        norm_chars.append(ch2)
        mapping.append(orig_idx)
    norm_text = "".join(norm_chars)
    needle2 = needle.replace("’", "'").replace("‘", "'").replace("“", '"').replace("”", '"')
    needle2 = needle2.replace("–", "-").replace("—", "-").replace("…", "...")
    needle2 = re.sub(r"\s+", " ", needle2).strip()
    pos = norm_text.find(needle2)
    if pos < 0:
        return None
    start = None
    for i in range(pos, len(mapping)):
        if mapping[i] is not None:
            start = mapping[i]
            break
    end = None
    for i in range(min(pos + len(needle2), len(mapping)) - 1, -1, -1):
        if mapping[i] is not None:
            end = mapping[i] + 1
            break
    if start is None or end is None:
        return None
    return start, end


def _colored_cell(cell, text, spans):
    """把一个单元格按 spans（(start,end,type) 已排序不重叠）拆成带色 run。"""
    from docx.shared import RGBColor
    cursor = 0
    first = True
    for start, end, atype in spans:
        if start > cursor:
            run = cell.paragraphs[0].add_run(text[cursor:start])
            _apply_run_fonts(run)
        run = cell.paragraphs[0].add_run(text[start:end])
        _apply_run_fonts(run)
        run.font.color.rgb = RGBColor.from_string(ANNOTATION_COLORS[atype])
        if atype in ("rare", "domain"):
            run.bold = True
        cursor = end
        first = False
    if cursor < len(text):
        run = cell.paragraphs[0].add_run(text[cursor:])
        _apply_run_fonts(run)
    if first and not text:
        cell.paragraphs[0].add_run("")


def pairs_to_word(pairs, annotations=None):
    """双语对照表 -> Word 表格。

    annotations: {seg: [{"type": "rare|domain|hard", "src_span": [s,e]|None,
                         "tgt_span": [s,e]|None, "note": str}]}
    """
    doc = Document()
    _apply_doc_fonts(doc)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "原文"
    table.rows[0].cells[1].text = "译文"
    annot = _normalize_annotations(annotations)
    for i, pair in enumerate(pairs):
        row = table.add_row().cells
        seg_annot = annot.get(i) or []
        src_spans = _compose_spans(
            [(it["src_span"][0], it["src_span"][1], it["type"])
             for it in seg_annot if it.get("src_span")], len(pair['source']))
        tgt_spans = _compose_spans(
            [(it["tgt_span"][0], it["tgt_span"][1], it["type"])
             for it in seg_annot if it.get("tgt_span")], len(pair['target']))
        _colored_cell(row[0], pair['source'], src_spans)
        _colored_cell(row[1], pair['target'], tgt_spans)
    # 图例（放表格后，避免挤占首行）
    p_legend = doc.add_paragraph()
    run = p_legend.add_run("图例：红色 = 生僻词/难词；黄色 = 专业名词（特殊译法）；"
                           "青绿色 = 翻译难点句（特别译法）。")
    _apply_run_fonts(run)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _add_formatted_runs(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 != 0:
            run.bold = True


def markdown_to_word(md_text, theory):
    doc = Document()
    _apply_doc_fonts(doc)
    md_text = re.sub(r'```markdown|```', '', md_text).strip()
    title = doc.add_heading(f'翻译实践报告：基于{theory}', 0)
    title.alignment = 1
    for line in md_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(('- ', '* ')):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, line[2:])
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Intense Quote')
            _add_formatted_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# ================= 任务持久化（真正的断点续传）=================
def _ensure_output_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def job_dir(job_id):
    return OUTPUT_DIR / job_id


def job_state_path(job_id):
    return job_dir(job_id) / "state.json"


def new_job_state(filename):
    state = {
        "filename": filename,
        "p1_done": False,
        "p2_done": False,
        "p3_done": False,
        "report_enabled": True,
        "paras": [],
        "pairs": [],
        "auto_terms": {},
        "findings": [],
        "review_stats": {
            "reviewed_segments": 0, "batches_reviewed": 0,
            "blocking": 0, "actionable": 0, "informational": 0, "review_failed": 0,
        },
        "tm_used_count": 0,
        "has_blocking": False,
        "p3_md": "",
        "p3_sections": [],
        "theory": "",
        "warnings": [],
        "annotations": {},
        "annotations_done": False,
        "annotations_done_offset": 0,
    }
    # 术语治理 / 交付门禁新增字段（默认值集中在 state_migration，保持单一来源）
    state.update(_state_migration._default_new_fields())
    return state


def load_job_state(job_id):
    p = job_state_path(job_id)
    if not p.is_file():
        return None
    try:
        raw = json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None
    return _state_migration.migrate_state(raw)


def save_job_state(job_id, state):
    """原子写入（先写临时文件再替换），避免中断写坏 state.json。"""
    d = job_dir(job_id)
    d.mkdir(parents=True, exist_ok=True)
    tmp = d / "state.json.tmp"
    tmp.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(d / "state.json")


def list_jobs():
    jobs = []
    _ensure_output_dir()
    for d in sorted(OUTPUT_DIR.iterdir()):
        sp = d / "state.json"
        if sp.is_file():
            try:
                s = json.loads(sp.read_text(encoding="utf-8"))
            except Exception:
                continue
            jobs.append({"job_id": d.name, "state": s})
    return jobs


def delete_job(job_id):
    d = job_dir(job_id)
    if d.exists():
        shutil.rmtree(d)


def file_job_id(file_bytes):
    """以文件内容哈希作为任务 ID：同一文件重传可自动续传，不同文件不会串状态。"""
    return hashlib.sha256(file_bytes).hexdigest()[:16]


def save_source(job_id, file_bytes):
    """留存源文件，刷新页面后即使不重新上传也能继续（如重做阶段一）。"""
    d = job_dir(job_id)
    d.mkdir(parents=True, exist_ok=True)
    (d / "source.bin").write_bytes(file_bytes)


def load_source(job_id):
    p = job_dir(job_id) / "source.bin"
    return p.read_bytes() if p.is_file() else None


def progress_label(state):
    if state.get("p1_done") and state.get("p2_done") and \
            (state.get("p3_done") or not state.get("report_enabled", True)):
        return "已完成"
    if state.get("p2_done"):
        return "报告生成中"
    if state.get("p1_done"):
        return "翻译中"
    return "待处理"


# ================= 术语审核状态（草稿 / 锁定 / 拒绝 / 冻结）=================
def save_glossary_draft(job_id, entries):
    """保存术语审核草稿（不冻结）。刷新/重启后从 TERMS_PREPARED 恢复。"""
    state = load_job_state(job_id)
    if state is None:
        return None
    norm = normalize_glossary(entries)
    state["glossary_draft"] = norm
    state["glossary"] = norm
    state["stage"] = "TERMS_PREPARED"
    save_job_state(job_id, state)
    return state


def set_glossary_entry_status(job_id, entry_ids, status):
    """批量修改术语状态（candidate/provisional/locked/rejected）。"""
    if status not in ("candidate", "provisional", "locked", "rejected"):
        raise ValueError(f"非法状态：{status}")
    state = load_job_state(job_id)
    if state is None:
        return None
    ids = set(entry_ids or [])
    for e in state.get("glossary") or []:
        if e.get("id") in ids:
            e["status"] = status
    save_job_state(job_id, state)
    return state


def freeze_glossary(job_id, entries=None, frozen_by="user"):
    """冻结术语表：生成新版本 + 确定性 glossary_hash。

    修改后再次冻结 -> 新版本追加到 glossary_versions，不悄悄覆盖旧冻结状态。
    """
    state = load_job_state(job_id)
    if state is None:
        return None
    norm = normalize_glossary(entries if entries is not None
                              else state.get("glossary") or [])
    source_hash = ""
    src = job_dir(job_id) / "source.bin"
    if src.is_file():
        source_hash = hashlib.sha256(src.read_bytes()).hexdigest()
    version = len(state.get("glossary_versions") or []) + 1
    frozen = {
        "version": version,
        "source_hash": source_hash,
        "entries": norm,
        "frozen_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "glossary_hash": _models.glossary_hash(norm),
        "frozen_by": frozen_by,
    }
    state["glossary"] = norm
    state["glossary_frozen"] = frozen
    state.setdefault("glossary_versions", []).append(frozen)
    state["stage"] = "GLOSSARY_FROZEN"
    if state.get("delivery_status") not in ("approved", "final"):
        state["delivery_status"] = "draft"
    save_job_state(job_id, state)
    return state


def save_document_profile(job_id, profile):
    """人工填写/修改文档画像后保存。"""
    state = load_job_state(job_id)
    if state is None:
        return None
    state["document_profile"] = _models.normalize_document_profile(profile)
    state["profile_done"] = True
    save_job_state(job_id, state)
    return state


def bypass_freeze(job_id, frozen_by="user"):
    """快速模式跳过人工冻结：允许以 provisional 术语直接翻译（记录审计标记）。"""
    state = load_job_state(job_id)
    if state is None:
        return None
    state["quality_bypass"] = True
    state.setdefault("human_actions", []).append({
        "action": "bypass_freeze",
        "note": "快速模式：跳过人工术语冻结，以 provisional 术语直接翻译",
        "timestamp": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "actor": frozen_by,
    })
    save_job_state(job_id, state)
    return state


# ================= 交付状态与人工处理记录 =================
def mark_findings_resolved(job_id, finding_ids, action, note="", actor="user"):
    """标记 findings 已人工处理，并重算交付状态。"""
    from mti_tool import delivery as _delivery
    state = load_job_state(job_id)
    if state is None:
        return None
    state, _marked = _delivery.mark_findings(state, finding_ids, action, note, actor)
    state["delivery_status"] = _delivery.compute_delivery_status(state)
    save_job_state(job_id, state)
    return state


def approve_delivery(job_id, note="", accept_blocking=False, actor="user"):
    """人工交付确认 -> final；有未解决 blocking 且不接受风险时拒绝。"""
    from mti_tool import delivery as _delivery
    state = load_job_state(job_id)
    if state is None:
        return None, False, ["任务不存在"]
    state, ok, errors = _delivery.approve_delivery(state, note, actor, accept_blocking)
    save_job_state(job_id, state)
    return state, ok, errors


def retranslate_segments(job_id, indexes, provider, api_key, model, target_lang,
                         style_rules="", glossary=None, on_status=None,
                         on_caption=None):
    """定点重译（抽取自 scripts/fix_segments.py 的能力）。"""
    from mti_tool import delivery as _delivery
    return _delivery.retranslate_segments(
        job_id, indexes, provider, api_key, model, target_lang,
        style_rules, glossary, on_status, on_caption)


def delivery_status_label(state):
    labels = {"draft": "草稿（draft）", "review_required": "待审（review_required）",
              "approved": "已批准（approved）", "final": "最终交付（final）"}
    return labels.get(state.get("delivery_status"), str(state.get("delivery_status")))


# ================= 阶段三：报告生成（Map-Reduce + 章节级断点）=================
def generate_mti_report(bilingual_pairs, termbase_dict, theory, provider, api_key,
                        model, state, job_id, on_status=None):
    """四段式报告生成；每个章节完成后立即落盘，中途失败可从已完成的章节续写。"""
    sample_texts = ""
    char_count = 0
    for pair in bilingual_pairs:
        chunk = f"【原文】{pair['source']}\n【译文】{pair['target']}\n\n"
        sample_texts += chunk
        char_count += len(chunk)
        if char_count > 8000:
            break

    term_str = "\n".join(f"{k} -> {v}" for k, v in termbase_dict.items()) if termbase_dict else "无术语提取"

    prompts = [
        (
            "一、 翻译项目概述与文本特征分析",
            f"请基于以下双语语料样本，撰写翻译实践报告的【第一部分】。\n"
            f"要求：详尽分析源文本的语言风格、专业领域、词法（如专有名词、长难句）与句法特点，"
            f"以及由此带来的总体翻译难点。字数要求 800-1000 字。严禁输出其他章节的内容。\n\n"
            f"语料样本：\n{sample_texts}"
        ),
        (
            "二、 术语管理与验证",
            f"请撰写翻译实践报告的【第二部分】。\n"
            f"请基于以下核心术语表，详细评估本次翻译中术语库的执行情况。"
            f"请至少选取 4 个核心术语，深度剖析其翻译策略（如直译、意译、增词、转换等）"
            f"及其对提升文本专业性的贡献。字数要求 800-1000 字。\n\n术语表：\n{term_str}"
        ),
        (
            f"三、 基于【{theory}】的案例分析",
            f"这是本报告的最核心章节。请基于以下双语语料样本和【{theory}】的理论框架，撰写报告的【第三部分】。\n"
            f"要求：精准抽取 4-5 个最具代表性的长难句或特殊表达案例。每个案例必须独立成段并包含：\n"
            f"1. 原译文对照\n2. 翻译难点深度剖析\n"
            f"3. 严谨的学理分析（明确指出具体的翻译技巧，并用【{theory}】的核心概念论证“为何如此翻译”）。\n"
            f"本部分字数要求不少于 1500 字，必须极具学术深度。\n\n语料样本：\n{sample_texts}"
        ),
        (
            "四、 翻译项目复盘与反思",
            f"请结合上述关于【{theory}】的翻译实践，撰写报告的【第四部分】。\n"
            f"要求：深刻总结机器翻译（MT）在此类文本中的局限性、本地化术语库强干预的实际效果，"
            f"以及作为译后编辑（MTPE）在双语能力和理论运用层面的收获。字数要求 600-800 字。"
        ),
    ]
    base_system_prompt = "你是一位拥有深厚学术背景的 MTI（翻译硕士）导师及资深学术期刊审稿人。" \
                         "请严格使用学术书面语，逻辑严密，杜绝任何 AI 常见的口语化或套话表达。"

    valid_titles = {title for title, _ in prompts}
    sections = [s for s in state.get("p3_sections", []) if s and s[0] in valid_titles]
    state["p3_sections"] = sections

    for idx, (section_title, user_prompt) in enumerate(prompts):
        if any(s[0] == section_title for s in sections):
            continue
        if on_status:
            on_status(f"【阶段三】正在深度撰写：{section_title} ({idx + 1}/4)...")
        last_err, success = None, False
        for _attempt in range(3):
            try:
                section_content = call_llm(provider, api_key, model, base_system_prompt,
                                           user_prompt, temperature=0.5)
                section_content = re.sub(r'^```markdown|```$', '',
                                         section_content.strip(), flags=re.MULTILINE)
                if not section_content.strip():
                    raise RuntimeError("模型返回空内容")
                sections.append([section_title, section_content.strip()])
                state["p3_sections"] = sections
                save_job_state(job_id, state)  # 章节级断点
                time.sleep(2)
                success = True
                break
            except Exception as e:
                last_err = e
                if is_rate_limited(e):
                    time.sleep(20)
                else:
                    break
        if not success:
            raise RuntimeError(f"报告章节「{section_title}」生成失败：{last_err}")

    return "".join(f"## {t}\n\n{c}\n\n---\n\n" for t, c in sections)


# ================= 主流程：单文档完整流水线 =================
def run_job_pipeline(job_id, filename, file_bytes, *, provider, api_key, model,
                     target_lang, auto_term, enable_report, translation_theory,
                     user_glossary=None, style_rules="", enable_review=True,
                     enable_annotate=True, mode="quick",
                     on_status=None, on_caption=None):
    """执行单个文档的完整流程；每个里程碑实时落盘，刷新/重启后均可继续。

    mode="quick"：自动术语作为 provisional，直接翻译（原行为）。
    mode="quality"：术语冻结（glossary_frozen）后才能开始翻译；
    未冻结时在 TERMS_PREPARED 阶段返回，由 UI 呈现术语审核面板。
    """
    _ensure_output_dir()
    base = new_job_state(filename)
    state = load_job_state(job_id) or base
    state = {**base, **state}  # 兼容旧版本状态缺字段
    state = _state_migration.migrate_state(state)
    state["report_enabled"] = bool(enable_report)
    warnings = state.setdefault("warnings", [])

    # 全部完成 -> 直接返回
    if state["p1_done"] and state["p2_done"] and (not enable_report or state["p3_done"]) \
            and (not enable_annotate or state.get("annotations_done")):
        state["stage"] = _state_migration.derive_stage(state)
        return state

    # ---------------- 阶段一：排版清洗 ----------------
    if not state["p1_done"]:
        if on_status:
            on_status("【阶段一】排版解析与段落重建（确定性提取）...")
        if file_bytes is None:
            file_bytes = load_source(job_id)
        if file_bytes is None:
            raise ValueError("缺少源文件，请重新上传后再继续")

        paragraphs = []
        if filename.lower().endswith(".pdf"):
            paragraphs = [clean_xml_chars(p) for p in extract_pdf_paragraphs(file_bytes)]
        elif filename.lower().endswith(".docx"):
            doc_word = Document(io.BytesIO(file_bytes))
            for p in doc_word.paragraphs:
                for sub_p in re.split(r'\n+', clean_xml_chars(p.text)):
                    t = sub_p.strip()
                    if len(t) > 1 and not _ORNAMENT_RE.match(t):
                        paragraphs.append(t)

        if not paragraphs:
            raise ValueError("未提取到有效文本（若为扫描版 PDF，请先做 OCR 生成文本层）")
        state["paras"] = paragraphs
        state["p1_done"] = True
        save_source(job_id, file_bytes)  # 留存源文件，刷新后无需重新上传
        save_job_state(job_id, state)

    # ---------------- 阶段 1.2：文档画像（分布式采样；失败仅警告，不阻断） ----------------
    if not state.get("profile_done"):
        if on_status:
            on_status("【阶段1.2】文档画像（分布式采样 + 结构化校验）...")
        from mti_tool.document_profile import profile_document
        profile, profile_warnings = profile_document(
            state["paras"], provider, api_key, model, target_lang)
        state["document_profile"] = profile
        state["profile_done"] = True
        for w in profile_warnings:
            if w not in warnings:
                warnings.append(w)
        if on_caption and profile:
            on_caption(f"✅ 文档画像完成：领域「{profile.get('domain') or '未知'}」"
                       f"· 文本类型「{profile.get('genre') or '未知'}」")
        elif on_caption:
            on_caption("⚠️ 文档画像失败，已跳过（可在 UI 中人工填写）。")
        save_job_state(job_id, state)

    # ---------------- 阶段 1.5：智能抽取术语 ----------------
    if auto_term and not state["auto_terms"]:
        if on_status:
            on_status("【阶段1.5】正在 AI 智能抽取全文核心术语...")
        if on_caption:
            on_caption("🤖 正在从全文分布式样本中提取专业术语...")
        from mti_tool.terminology import extract_auto_terms_v2
        entries, extract_warnings = extract_auto_terms_v2(
            state["paras"], target_lang, provider, api_key, model,
            document_profile=state.get("document_profile"))
        state["auto_term_entries"] = entries
        state["auto_terms"] = {e["source"]: e["target"] for e in entries}
        if entries:
            if on_caption:
                on_caption(f"✅ 成功提取 {len(entries)} 个候选术语（全部出现位置已记录）")
        else:
            msg = "术语抽取失败（限流或返回格式异常），已跳过该步骤；可稍后点击“继续处理”重试。"
            if msg not in warnings and msg not in extract_warnings:
                warnings.append(msg)
        for w in extract_warnings:
            if w not in warnings:
                warnings.append(w)
        save_job_state(job_id, state)
    legacy_auto = [{"source": k, "target": v, "behavior": "translate",
                    "status": "provisional"}
                   for k, v in (state["auto_terms"] or {}).items()]
    auto_entries = normalize_glossary(state.get("auto_term_entries") or legacy_auto)
    user_entries = normalize_glossary(list(user_glossary or []))
    if mode == "quick":
        for e in auto_entries:
            if e["status"] == "candidate":
                e["status"] = "provisional"
    working = normalize_glossary(state.get("glossary") or [])
    if not working:
        working = normalize_glossary(user_entries + auto_entries)
    else:
        # 新上传/新抽取的术语若不在已保存审核表中，追加（不覆盖人工审核结果）
        known = {e["source"].casefold() for e in working}
        for e in user_entries + auto_entries:
            if e["source"].casefold() not in known:
                working.append(e)
    state["glossary"] = working
    glossary = working
    final_termbase = glossary_to_terms(glossary)

    # ---------------- 高质量模式门禁：术语冻结后才能开始翻译 ----------------
    if mode == "quality":
        state["quality_mode"] = True
        if not state.get("glossary_frozen") and not state.get("quality_bypass"):
            if on_status:
                on_status("⏸ 高质量模式：等待人工术语审核与冻结（术语面板）...")
            msg = ("高质量模式：术语尚未冻结，翻译未开始。"
                   "请在“术语准备与审核”面板完成冻结后继续。")
            if msg not in warnings:
                warnings.append(msg)
            state["stage"] = _state_migration.derive_stage(state)
            save_job_state(job_id, state)
            return state

    # ---------------- 阶段二：双语翻译（批次 + 确定性检查 + 独立审校 + 翻译记忆）----------------
    if not state["p2_done"]:
        if on_status:
            on_status("【阶段二】双语翻译与术语严格注入（批次翻译 + 确定性检查 + 独立审校）...")
        translate_stage(state, job_id, glossary, provider, api_key, model, target_lang,
                        style_rules, enable_review,
                        document_profile=state.get("document_profile"),
                        on_status=on_status, on_caption=on_caption)
        state["p2_done"] = True
        state["delivery_status"] = "review_required" if state.get("has_blocking") \
            else "draft"
        save_job_state(job_id, state)

    # ---------------- 阶段 2.5：三色自动标注 ----------------
    if enable_annotate and state["p2_done"] and not state.get("annotations_done"):
        if on_status:
            on_status("【阶段 2.5】自动标注学习重点（红=生僻词 / 黄=专业名词 / 青绿=难点句）...")
        annotate_stage(state, job_id, glossary, provider, api_key, model, target_lang,
                       on_caption=on_caption)

    # ---------------- 阶段三：报告生成 ----------------
    if enable_report and not state["p3_done"]:
        if on_status:
            on_status(f"【阶段三】基于《{translation_theory}》生成报告...")
        report_md = generate_mti_report(state["pairs"], final_termbase, translation_theory,
                                        provider, api_key, model, state, job_id,
                                        on_status=on_status)
        if not report_md.strip():
            raise RuntimeError("报告内容为空，请点击“继续处理”重试")
        state["p3_md"] = report_md
        state["theory"] = translation_theory
        state["p3_done"] = True
        save_job_state(job_id, state)

    state["stage"] = _state_migration.derive_stage(state)
    return state
