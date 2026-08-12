#!/usr/bin/env python3
"""Assemble the validated v6 thesis into the 2026 college DOCX format.

The reference report template supplies the Word theme and style package. The
content is rebuilt from the frozen v6 thesis body and audited project state;
historical translation evidence is read-only.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUT_DIR = (
    ROOT / "eval/academic-quality/ec100d8686d3891e/thesis-closeout-v6"
)
DEFAULT_TEMPLATE = Path(
    "/Users/xueyang/Documents/毕业论文/论文要求/3-MTI论文写作规范/"
    "1-MTI翻译实践报告模板（也适用于翻译项目报告）-20251121.docx"
)

TITLE_ZH = "《当天空属于我们时》英汉翻译实践报告\n——叙事关系、文化信息与语篇回指的证据化分析"
TITLE_EN = (
    "A Report on the E-C Translation of When the Sky Was Ours:\n"
    "An Evidence-Based Analysis of Narrative Relations, Cultural Information,\n"
    "and Discourse Reference"
)

# Word can map the generic Chinese names on the author's machine, but the
# headless renderer cannot. These installed macOS faces are metrically close
# to the college's Song/Hei requirements and render consistently in both.
BODY_CN_FONT = "Songti SC"
HEADING_CN_FONT = "Heiti SC"
COVER_CN_FONT = "FZXiaoBiaoSong-B05S"

ABSTRACT_ZH = (
    "本报告以文学性军事回忆录《当天空属于我们时》的英汉翻译实践为对象，重点考察人物关系意象、"
    "文化专名和语篇回指在汉译中的处理。源文兼具飞行与军旅信息、第一人称回忆、口语对话和跨文化"
    "指称等特征。项目以段落为基本处理单位，保存源文、初译、终译和审校记录，并对273个源译段开展"
    "全量复核；其中188段使用了翻译记忆，34条历史可处理问题均已完成处置。为避免把技术性错位包装"
    "成翻译策略，报告先执行真实初译—终译修订准入审查，再从22项合格修订中选取三项证据互补的"
    "核心案例。分析表明：汉语句法压缩可以与人物关系意象的保留并行；文化对象核验和人称回指恢复"
    "需要分别处理；英语计词表达转入汉语时，应避免把词数机械转换为字数，整句回指能够消除目标语"
    "内部矛盾。报告同时以翻译质量评价、文学翻译衔接和叙事视角研究约束结论强度。研究说明，基于"
    "可核验文本变化的审校流程能够提高案例分析的可追溯性，但单一项目、缺少译者同期说明和没有"
    "中文读者实验限制了结论的外推范围。"
)
KEYWORDS_ZH = "军事回忆录；文学翻译；叙事关系；文化专名；语篇回指；翻译质量评价"

ABSTRACT_EN = (
    "This report examines the English-Chinese translation of the literary military memoir "
    "When the Sky Was Ours, with particular attention to relational imagery, cultural names, "
    "and discourse reference. The source text combines aviation and military information with "
    "first-person recollection, conversational language, and cross-cultural references. The "
    "project retained the source text, initial translation, final translation, and review record "
    "for each paragraph and reviewed all 273 source-target segments. Translation memory was used "
    "for 188 segments, and all 34 historically actionable review issues were resolved. To prevent "
    "technical misalignment from being misrepresented as a translation strategy, the report first "
    "applied an eligibility check for genuine initial-to-final revisions and then selected three "
    "complementary cases from 22 eligible revisions. The analysis shows that syntactic compression "
    "in Chinese can coexist with the retention of relational imagery; cultural-object verification "
    "and the restoration of personal reference require separate treatment; and English word-count "
    "expressions should not be mechanically converted into Chinese character counts, because a "
    "whole-utterance reference can remove an internal inconsistency in the target text. Research on "
    "translation quality assessment, cohesion in literary translation, and narrative viewpoint is "
    "used to constrain rather than replace the textual analysis. The study suggests that an "
    "evidence-grounded review process improves the traceability of case analysis, while the single "
    "project, absence of contemporaneous translator accounts, and lack of Chinese reader testing "
    "limit generalisation."
)
KEYWORDS_EN = (
    "military memoir; literary translation; narrative relations; cultural names; "
    "discourse reference; translation quality assessment"
)

REFERENCES = [
    "House, J. (2001). Translation quality assessment: Linguistic description versus social "
    "evaluation. Meta, 46(2), 243–257. https://doi.org/10.7202/003141ar",
    "Károly, K., Karádi, G., Olgyay-Fekete, J., & Sulyok, K. (2022). A szövegkohézió "
    "újrateremtése műfordításban: esettanulmány Salinger The Catcher in the Rye című művének "
    "két magyar fordításáról. Fordítástudomány, 24(2), 5–40. "
    "https://doi.org/10.35924/fordtud.24.2.1",
    "Károly, K., Csiborné Horváth, A., Engel, I., & Van Waarden, F. (2022). A makrostruktúra "
    "újrateremtése műfordításban: esettanulmány Salinger The Catcher in the Rye című művének "
    "két magyar fordításáról. Fordítástudomány, 24(2), 41–81. "
    "https://doi.org/10.35924/fordtud.24.2.2",
    "Eekhof, L. S., van Krieken, K., & Sanders, J. (2020). VPIP: A lexical identification "
    "procedure for perceptual, cognitive, and emotional viewpoint in narrative discourse. "
    "Open Library of Humanities, 6(1), Article 18, 1–38. https://doi.org/10.16995/olh.483",
    "van Krieken, K. (2018). Ambiguous perspective in narrative discourse: Effects of viewpoint "
    "markers and verb tense on readers’ interpretation of represented perceptions. Discourse "
    "Processes, 55(8), 771–786. https://doi.org/10.1080/0163853X.2017.1381540",
    "Al Herz, K. H. H. (2016). Narrative point of view in translation: A systemic functional "
    "analysis of the Arabic translations of J. M. Coetzee’s Waiting for the Barbarians "
    "[Doctoral dissertation, University of Leeds]. White Rose eTheses Online. "
    "https://etheses.whiterose.ac.uk/id/eprint/17870/",
]

CASE_LABELS = {"0139": "案例一", "0233": "案例二", "0272": "案例三"}
TOC_ENTRIES = [
    (1, "1 引言"),
    (2, "1.1 研究背景及意义"), (2, "1.2 研究问题"), (2, "1.3 报告结构"),
    (1, "2 翻译项目概述"),
    (2, "2.1 项目简介"), (2, "2.2 翻译流程"),
    (3, "2.2.1 译前准备"), (3, "2.2.2 翻译过程"), (3, "2.2.3 译后管理"),
    (1, "3 翻译项目案例分析"),
    (2, "3.1 源语文本的类型与特征"), (2, "3.2 翻译难点"),
    (2, "3.3 翻译策略与解决方案"),
    (3, "3.3.1 案例一：隐喻保留与叙事节奏"),
    (3, "3.3.2 案例二：文化对象识别与明确回指"),
    (3, "3.3.3 案例三：元语言计量与整句回指"),
    (3, "3.3.4 跨案例综合"),
    (1, "4 总结与反思"),
    (2, "4.1 研究问题回应"), (2, "4.2 实践经验与可迁移方法"),
    (2, "4.3 局限与改进方向"),
    (1, "参考文献"), (1, "致  谢"),
    (1, "附录一 核心案例原文、初译与终译"),
    (1, "附录二 翻译实践原文与终译"),
]


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _set_run_font(run, *, size=12, bold=None, cn=BODY_CN_FONT, latin="Times New Roman"):
    # LibreOffice on macOS does not consistently honor only w:eastAsia. Give
    # CJK-bearing runs the installed Chinese family across all font slots;
    # English-only runs keep Times New Roman.
    primary = cn if re.search(r"[\u3400-\u9fff]", run.text or "") else latin
    run.font.name = primary
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), primary)
    rfonts.set(qn("w:hAnsi"), primary)
    rfonts.set(qn("w:eastAsia"), cn)


def _set_style_font(style, *, size, cn, latin, bold=None):
    style.font.name = latin
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cn)


def _clear_document_body(doc: Document) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, size=12, cn=BODY_CN_FONT, latin="Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    for name, size, align in (
        ("Heading 1", 15, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 14, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 14, WD_ALIGN_PARAGRAPH.LEFT),
    ):
        style = doc.styles[name]
        _set_style_font(style, size=size, cn=HEADING_CN_FONT, latin="Times New Roman", bold=True)
        style.paragraph_format.alignment = align
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.space_before = Pt(8 if name != "Heading 1" else 0)
        style.paragraph_format.space_after = Pt(6)


def _configure_section(section, *, header=False, footer=False, page_format=None, start=1):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.different_first_page_header_footer = False
    _clear_container(section.header)
    _clear_container(section.footer)
    if header:
        p = section.header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("南京航空航天大学翻译硕士学位论文")
        _set_run_font(run, size=9, cn=BODY_CN_FONT, latin="Times New Roman")
        ppr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), "000000")
        borders.append(bottom)
        ppr.append(borders)
    if footer:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        _set_run_font(run, size=10)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run._r.append(fld)
    if page_format:
        sect_pr = section._sectPr
        pg_num = sect_pr.find(qn("w:pgNumType"))
        if pg_num is None:
            pg_num = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num)
        pg_num.set(qn("w:fmt"), page_format)
        pg_num.set(qn("w:start"), str(start))


def _clear_container(container) -> None:
    for p in container.paragraphs:
        for child in list(p._p):
            p._p.remove(child)


def _remove_unreferenced_headers_and_footers(docx_path: Path) -> None:
    """Drop unused template header/footer parts that python-docx leaves behind."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    pr_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    ns = {"w": w_ns, "r": r_ns, "pr": pr_ns, "ct": ct_ns}
    from lxml import etree

    with zipfile.ZipFile(docx_path, "r") as source:
        contents = {name: source.read(name) for name in source.namelist()}

    document = etree.fromstring(contents["word/document.xml"])
    rels = etree.fromstring(contents["word/_rels/document.xml.rels"])
    rel_by_id = {
        rel.get("Id"): rel
        for rel in rels.xpath(".//pr:Relationship", namespaces=ns)
    }
    referenced_ids = {
        node.get(f"{{{r_ns}}}id")
        for node in document.xpath(
            ".//w:sectPr/w:headerReference|.//w:sectPr/w:footerReference",
            namespaces=ns,
        )
    }
    retained_targets = {
        rel_by_id[rel_id].get("Target")
        for rel_id in referenced_ids
        if rel_id in rel_by_id
    }
    removed_targets: set[str] = set()
    for rel in list(rels):
        rel_type = rel.get("Type", "")
        target = rel.get("Target", "")
        if rel_type.endswith(("/header", "/footer")) and target not in retained_targets:
            removed_targets.add(target)
            rels.remove(rel)
    contents["word/_rels/document.xml.rels"] = etree.tostring(
        rels, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    content_types = etree.fromstring(contents["[Content_Types].xml"])
    removed_parts = {f"/word/{target}" for target in removed_targets}
    for override in list(content_types):
        if override.get("PartName") in removed_parts:
            content_types.remove(override)
    contents["[Content_Types].xml"] = etree.tostring(
        content_types, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    for target in removed_targets:
        contents.pop(f"word/{target}", None)
        stem = Path(target).name
        contents.pop(f"word/_rels/{stem}.rels", None)

    with tempfile.NamedTemporaryFile(
        dir=docx_path.parent,
        prefix=f".{docx_path.name}.",
        suffix=".cleaning",
        delete=False,
    ) as temp_handle:
        temp_path = Path(temp_handle.name)
    try:
        with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as out:
            for name, data in contents.items():
                out.writestr(name, data)
        temp_path.replace(docx_path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def _new_section(doc: Document, **kwargs):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(section, **kwargs)
    return section


def _add_spacer(doc: Document, points=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)
    return p


def _add_centered(doc: Document, text: str, *, size=14, bold=False, cn=BODY_CN_FONT, latin="Times New Roman", after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    for idx, line in enumerate(text.split("\n")):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        _set_run_font(run, size=size, bold=bold, cn=cn, latin=latin)
    return p


def _add_body(
    doc: Document,
    text: str,
    *,
    first_indent=True,
    size=12,
    cn=BODY_CN_FONT,
    latin="Times New Roman",
    line_spacing=1.5,
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.first_line_indent = Pt(24) if first_indent else Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.widow_control = True
    run = p.add_run(text)
    _set_run_font(run, size=size, cn=cn, latin=latin)
    return p


def _add_labelled(doc: Document, label: str, text: str, *, size=11, shade=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.right_indent = Pt(12)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    if shade:
        ppr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), shade)
        ppr.append(shd)
    lr = p.add_run(label)
    _set_run_font(lr, size=size, bold=True, cn=HEADING_CN_FONT)
    tr = p.add_run(text)
    _set_run_font(tr, size=size)
    return p


def _add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_heading(doc: Document, text: str, level: int, *, page_break_before=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.page_break_before = page_break_before
    run = p.add_run(text)
    _set_run_font(run, size=15 if level == 1 else 14, bold=True, cn=HEADING_CN_FONT)
    return p


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _format_table_cell(cell, text: str, *, size=9.5, bold=False, cn=BODY_CN_FONT, latin="Times New Roman"):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, cn=cn, latin=latin)


def _add_metadata_table(doc: Document):
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.0)
    table.columns[1].width = Cm(8.5)
    labels = ["研究生姓名", "专业类别", "专业领域", "指导教师", "行业导师"]
    values = ["[待填写：姓名]", "翻译硕士", "英语笔译", "[待填写：姓名及职称]", "[选填]"]
    _set_repeat_table_header(table.rows[0])
    for row, label, value in zip(table.rows, labels, values):
        _format_table_cell(row.cells[0], label, size=14, cn=BODY_CN_FONT)
        _format_table_cell(row.cells[1], value, size=14, cn=BODY_CN_FONT)
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "nil")
                borders.append(el)
            tc_pr.append(borders)


def _clean_text(text: str) -> str:
    text = re.sub(r"<!--.*?-->", "", text)
    text = text.replace("0139要求", "案例一要求")
    text = text.replace("0233要求", "案例二要求")
    text = text.replace("0272要求", "案例三要求")
    text = text.replace("前两例有审校问题和技术性修复说明", "案例一与案例二有审校问题和技术性修复说明")
    return text.strip()


def _clean_heading(text: str) -> str:
    text = _clean_text(text)
    replacements = {
        "3.3.1 真实修订案例0139：": "3.3.1 案例一：",
        "3.3.2 真实修订案例0233：": "3.3.2 案例二：",
        "3.3.3 真实修订案例0272：": "3.3.3 案例三：",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def _add_thesis_body(doc: Document, markdown: str):
    paragraphs: list[str] = []
    buffer: list[str] = []

    def flush():
        if buffer:
            paragraphs.append(" ".join(x.strip() for x in buffer).strip())
            buffer.clear()

    for line in markdown.splitlines():
        if not line.strip():
            flush()
            continue
        if line.startswith(("## ", "### ", "#### ", "[seg-", "> [")):
            flush()
            paragraphs.append(line)
        else:
            buffer.append(line)
    flush()

    chapter_seen = 0
    for raw in paragraphs:
        if raw.startswith("[seg-"):
            continue
        if raw.startswith("## "):
            chapter_seen += 1
            _add_heading(
                doc,
                _clean_heading(raw[3:]),
                1,
                page_break_before=chapter_seen > 1,
            )
            continue
        if raw.startswith("### "):
            _add_heading(doc, _clean_heading(raw[4:]), 2)
            continue
        if raw.startswith("#### "):
            _add_heading(doc, _clean_heading(raw[5:]), 3)
            continue
        if raw.startswith("> [SOURCE "):
            _add_labelled(doc, "原文：", raw.split("]: ", 1)[1], shade="F2F2F2")
            continue
        if raw.startswith("> [INITIAL "):
            _add_labelled(doc, "初译：", raw.split("]: ", 1)[1], shade="F2F2F2")
            continue
        if raw.startswith("> [TARGET "):
            _add_labelled(doc, "终译：", raw.split("]: ", 1)[1], shade="F2F2F2")
            continue
        cleaned = _clean_text(raw)
        if cleaned:
            _add_body(doc, cleaned)


def _add_toc(doc: Document, page_map: dict[str, str]):
    _add_heading(doc, "目  录", 1)
    for level, title in TOC_ENTRIES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm({1: 0, 2: 0.75, 3: 1.5}[level])
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.0), WD_ALIGN_PARAGRAPH.RIGHT, leader=2
        )
        r = p.add_run(f"{title}\t{page_map.get(title, '—')}")
        _set_run_font(r, size=10.5, cn=BODY_CN_FONT, latin="Times New Roman")


def _add_references(doc: Document):
    _add_heading(doc, "参考文献", 1, page_break_before=True)
    for idx, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.first_line_indent = Pt(-21)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"[{idx}] {ref}")
        _set_run_font(run, size=10.5, cn=BODY_CN_FONT, latin="Times New Roman")


def _add_acknowledgements(doc: Document):
    _add_heading(doc, "致  谢", 1, page_break_before=True)
    _add_body(
        doc,
        "[待作者本人根据真实指导、协助与支持情况补写。致谢属于个人事实陈述，送审稿不代拟姓名、经历或关系。]",
    )


def _add_core_case_appendix(doc: Document, evidence: dict):
    _add_heading(doc, "附录一 核心案例原文、初译与终译", 1, page_break_before=True)
    segments = {x["segment_index"]: x for x in evidence["project_evidence"]["segments"]}
    for index, key in enumerate((139, 233, 272), 1):
        if index > 1:
            _add_spacer(doc, 6)
        _add_heading(doc, f"{CASE_LABELS[f'{key:04d}']}：修订证据", 2)
        item = segments[key]
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(2.2)
        table.columns[1].width = Cm(13.0)
        _format_table_cell(table.rows[0].cells[0], "证据类型", size=10, bold=True, cn=HEADING_CN_FONT)
        _format_table_cell(table.rows[0].cells[1], "文本", size=10, bold=True, cn=HEADING_CN_FONT)
        for cell in table.rows[0].cells:
            _set_cell_shading(cell, "D9E2F3")
        rows = (
            ("原文", item["source"]),
            ("初译", item["initial_target"]),
            ("终译", item["final_target"]),
        )
        _set_repeat_table_header(table.rows[0])
        for row, (label, value) in zip(table.rows[1:], rows):
            _format_table_cell(row.cells[0], label, size=10.5, bold=True, cn=HEADING_CN_FONT)
            _set_cell_shading(row.cells[0], "E7E6E6")
            _format_table_cell(row.cells[1], value, size=10.5)


def _add_full_translation_appendix(doc: Document, pairs: list[dict]):
    _add_heading(doc, "附录二 翻译实践原文与终译", 1, page_break_before=True)
    note = (
        "本附录按项目段落顺序列出英语源文与审校后的中文终译，用于核对报告所述文本范围。"
        "初译只在附录一的核心案例中呈现。"
    )
    _add_body(doc, note)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(7.6)
    table.columns[1].width = Cm(7.6)
    _format_table_cell(table.rows[0].cells[0], "Source Text", size=10, bold=True)
    _format_table_cell(table.rows[0].cells[1], "终译", size=10, bold=True, cn=HEADING_CN_FONT)
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, "D9E2F3")
    _set_repeat_table_header(table.rows[0])
    for pair in pairs:
        row = table.add_row()
        _format_table_cell(row.cells[0], pair["source"], size=9.5)
        _format_table_cell(row.cells[1], pair["target"], size=9.5)


def _write_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("中图分类号：H315.9                           论文编号：[待填写]")
    _set_run_font(r, size=12)
    p = doc.add_paragraph()
    r = p.add_run("学科分类号：055101")
    _set_run_font(r, size=12)
    _add_spacer(doc, 6)
    _add_centered(doc, "专业学位硕士学位论文", size=26, cn=COVER_CN_FONT, after=8)
    _add_centered(doc, TITLE_ZH, size=18, cn=COVER_CN_FONT, after=6)
    _add_metadata_table(doc)
    _add_spacer(doc, 2)
    _add_centered(doc, "南京航空航天大学", size=18, cn="FangSong_GB2312")
    _add_centered(doc, "研究生院 外国语学院    [待填写：提交年月]", size=11)


def _write_english_title_page(doc: Document):
    _add_spacer(doc, 8)
    _add_centered(doc, "Nanjing University of Aeronautics and Astronautics", size=14)
    _add_centered(doc, "The Graduate School", size=14)
    _add_centered(doc, "College of Foreign Languages", size=14, after=30)
    _add_centered(doc, TITLE_EN, size=16, bold=True, after=10)
    for text in (
        "A Thesis in", "Translation and Interpreting", "by", "[Name to be completed]",
        "Advised by", "[Supervisor to be completed]",
    ):
        _add_centered(doc, text, size=11, after=1)
    _add_spacer(doc, 2)
    for text in (
        "Submitted in Partial Fulfillment", "of the Requirements", "for the Degree of",
        "Master of Translation and Interpreting", "[Submission date to be completed]",
    ):
        _add_centered(doc, text, size=11, after=0)


def _write_declarations(doc: Document):
    _add_centered(doc, "独创性声明", size=20, cn=HEADING_CN_FONT, after=14)
    _add_body(
        doc,
        "本人声明所呈交的硕士学位论文是本人在导师指导下进行的研究工作及取得的研究成果。除了文中"
        "特别加以标注和致谢的地方外，论文中不包含其他人已经发表或撰写过的研究成果，也不包含为获得"
        "南京航空航天大学或其他教育机构的学位或证书而使用过的材料。",
        size=12,
    )
    _add_spacer(doc, 4)
    _add_body(doc, "研究生签名：________________    日期：________________", first_indent=False, size=12)
    _add_page_break(doc)
    _add_centered(doc, "使用授权声明", size=20, cn=HEADING_CN_FONT, after=14)
    _add_body(
        doc,
        "本人完全了解南京航空航天大学有关保留、使用学位论文的规定，即学校有权保留并向国家有关部门"
        "或机构送交论文的复印件和电子版，允许论文被查阅和借阅。本人授权南京航空航天大学可以将本学位"
        "论文的全部内容编入有关数据库进行检索，可以采用影印、缩印或扫描等复制手段保存、汇编学位论文。",
        size=12,
    )
    _add_spacer(doc, 4)
    _add_body(
        doc,
        "研究生签名：____________    导师签名：____________    日期：____________",
        first_indent=False,
        size=12,
    )


def _write_front_matter(doc: Document, page_map: dict[str, str]):
    _add_heading(doc, "摘  要", 1)
    _add_body(doc, ABSTRACT_ZH, size=10.5, line_spacing=1.3)
    _add_spacer(doc, 4)
    p = _add_body(doc, "", first_indent=False, size=10.5, line_spacing=1.3)
    lr = p.add_run("关键词：")
    _set_run_font(lr, size=10.5, bold=True, cn=HEADING_CN_FONT)
    kr = p.add_run(KEYWORDS_ZH)
    _set_run_font(kr, size=10.5)

    _add_page_break(doc)
    _add_heading(doc, "ABSTRACT", 1)
    _add_body(doc, ABSTRACT_EN, size=10.5, cn=BODY_CN_FONT, latin="Times New Roman")
    _add_spacer(doc, 4)
    p = _add_body(doc, "", first_indent=False, size=10.5)
    lr = p.add_run("Keywords: ")
    _set_run_font(lr, size=10.5, bold=True, latin="Times New Roman")
    kr = p.add_run(KEYWORDS_EN)
    _set_run_font(kr, size=10.5, latin="Times New Roman")

    _add_page_break(doc)
    _add_toc(doc, page_map)


def _update_closeout_state(out_dir: Path, output_docx: Path, page_map: dict[str, str]) -> None:
    state_path = out_dir / "thesis-closeout-state.json"
    state = json.loads(state_path.read_text())
    state["status"] = "review_package_assembled_supervisor_review_pending"
    for stage in state["stages"]:
        if stage["stage"] == 8:
            stage["status"] = "review_draft_assembled"
    state["document_assembly"] = {
        "status": "review_draft_assembled",
        "format": "docx",
        "body_language": "zh-CN",
        "output": output_docx.name,
        "template": DEFAULT_TEMPLATE.name,
        "template_sha256": _sha256(DEFAULT_TEMPLATE),
        "toc_page_map_applied": bool(page_map),
        "pending_before_submission": [
            "supervisor_semantic_review",
            "author_and_supervisor_metadata",
            "personal_acknowledgements",
            "signature_and_submission_dates",
        ],
    }
    state_path.write_text(json.dumps(state, ensure_ascii=False, indent=2) + "\n")


def build(template: Path, out_dir: Path, output_docx: Path, page_map: dict[str, str]) -> dict:
    body_path = out_dir / "thesis-body-v6.md"
    evidence_path = out_dir / "academic-evidence-final.json"
    state_path = out_dir / "translation-state-after-review.json"
    for required in (template, body_path, evidence_path, state_path):
        if not required.is_file():
            raise FileNotFoundError(required)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(template, output_docx)
    doc = Document(output_docx)
    _clear_document_body(doc)
    _configure_styles(doc)
    doc.core_properties.title = TITLE_ZH.replace("\n", "")
    doc.core_properties.subject = "2026年MTI中文翻译实践报告送审稿"
    doc.core_properties.author = ""
    doc.core_properties.comments = "基于学院模板装配；导师人工复核与个人信息待补。"

    _configure_section(doc.sections[0], header=False, footer=False)
    _write_cover(doc)
    _new_section(doc, header=False, footer=False)
    _write_english_title_page(doc)
    _new_section(doc, header=False, footer=False)
    _write_declarations(doc)
    _new_section(doc, header=True, footer=True, page_format="upperRoman", start=1)
    _write_front_matter(doc, page_map)
    _new_section(doc, header=True, footer=True, page_format="decimal", start=1)
    _add_thesis_body(doc, body_path.read_text())
    _add_references(doc)
    _add_acknowledgements(doc)

    evidence = json.loads(evidence_path.read_text())
    state = json.loads(state_path.read_text())
    _add_core_case_appendix(doc, evidence)
    _add_full_translation_appendix(doc, state["pairs"])
    doc.save(output_docx)
    _remove_unreferenced_headers_and_footers(output_docx)
    _update_closeout_state(out_dir, output_docx, page_map)

    manifest = {
        "schema_version": "thesis-document-assembly-v1",
        "status": "review_draft_assembled",
        "template": str(template),
        "template_sha256": _sha256(template),
        "output": str(output_docx),
        "output_sha256": _sha256(output_docx),
        "body_source": str(body_path),
        "body_source_sha256": _sha256(body_path),
        "sections": 5,
        "chapters": 4,
        "core_cases": 3,
        "references": len(REFERENCES),
        "appendix_translation_pairs": len(state["pairs"]),
        "toc_page_map": page_map,
        "supervisor_review_status": "pending_supervisor_review",
    }
    (out_dir / "thesis-document-assembly-manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n"
    )
    return manifest


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUT_DIR)
    parser.add_argument("--output", type=Path, default=None)
    parser.add_argument("--page-map", type=Path, default=None)
    args = parser.parse_args()
    output = args.output or (args.out_dir / "thesis-review-v6.docx")
    page_map = json.loads(args.page_map.read_text()) if args.page_map else {}
    manifest = build(args.template, args.out_dir, output, page_map)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
