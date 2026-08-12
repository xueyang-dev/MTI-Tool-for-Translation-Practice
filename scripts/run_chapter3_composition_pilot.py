"""Compose and evaluate an isolated Chapter 3 pilot from frozen case evidence."""
from __future__ import annotations

import argparse
import copy
import hashlib
import json
import os
import re
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import core
from mti_tool import academic_evidence, academic_quality, academic_validator
from mti_tool import academic_writer, case_analysis


JOB = "ec100d8686d3891e"
AUTHENTIC = [f"seg-{JOB}-0209", f"seg-{JOB}-0272"]
SYNTHETIC = "SC-0141"
PILOT_VERSION = "chapter3-composition-pilot-v1"
REVIEW_VERSION = "chapter3-supervisor-review-v1"


def _load(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _write_json(path: Path, value: dict[str, Any]) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n",
                    encoding="utf-8")


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _has_invalid_0272_count_claim(text: str) -> bool:
    """Reject generated prose that miscounts the visible English or Chinese quote."""
    patterns = (
        r"源文为五个英文词",
        r"英文引语.{0,16}(?:恰好|正好)为?五个?词",
        r"(?:恰好|正好)五字.{0,16}你不会有战争",
        r"你不会有战争.{0,16}(?:恰好|正好)五字",
    )
    for sentence in re.split(r"[。；\n]", text):
        for pattern in patterns:
            match = re.search(pattern, sentence)
            if not match:
                continue
            prefix = sentence[:match.start()]
            if not re.search(r"(?:不得|不能|不应|并非|不是).{0,36}$", prefix):
                return True
    return False


def _stamp(value: dict[str, Any]) -> dict[str, Any]:
    value["content_hash"] = academic_evidence.stable_hash(
        {k: v for k, v in value.items() if k != "content_hash"})
    return value


def _json_call(api_key: str, system: str, payload: dict[str, Any],
               provider: str, model: str) -> dict[str, Any]:
    for attempt in range(2):
        prompt = system if not attempt else system + "\n上次输出无效；只输出合法 JSON。"
        raw = core.call_llm(provider, api_key, model, prompt,
                            json.dumps(payload, ensure_ascii=False), temperature=0.1)
        candidate = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw.strip(),
                           flags=re.DOTALL)
        try:
            value = json.loads(candidate)
        except json.JSONDecodeError:
            continue
        if isinstance(value, dict):
            return value
    raise RuntimeError("review model did not return a JSON object")


def _research_model() -> dict[str, Any]:
    return _stamp({
        "schema_version": academic_writer.VERSIONS["research_model_version"],
        "research_topic": "真实修订证据与受控合成对比在翻译案例分析中的互补作用",
        "research_questions": [
            {"rq_id": "RQ1", "question": "两项真实初译至终译记录分别揭示了哪些可观察的修订问题与文本变化？", "provenance": "pilot_inferred"},
            {"rq_id": "RQ2", "question": "在不补造第三项历史修订的前提下，受控合成对比能否补充一种独特且边界清楚的翻译问题分析？", "provenance": "pilot_inferred"},
        ],
        "theoretical_framework": [],
        "method": "两项历史修订的证据约束分析，加一项补充性受控合成对比",
        "analysis_dimensions": ["源译对应关系", "话语指称", "语用力度", "证据边界"],
        "expected_contribution": [
            "区分历史修订事实、作者分析与合成对比证据",
            "检验合成案例能否在不冒充翻译历史的条件下补足机制层分析",
        ],
        "writing_style": "规范、克制的中文 MTI 学术书面语",
        "report_requirements": "翻译实践报告第三章",
        "target_words": 2600,
        "settings_provenance": {"research_topic": "pilot_inferred", "method": "pilot_inferred"},
    })


def _argument_plan() -> dict[str, Any]:
    claims = [
        {
            "claim_id": "C1",
            "claim": "0209 的可证实修订是把与源文无对应关系的中文片段替换为题名文本，从而恢复源译段对应关系；现有记录不能证明保留英文题名是有意的最终翻译策略。",
            "research_question": "RQ1", "project_evidence": [AUTHENTIC[0]],
            "synthetic_contrast_evidence": [], "literature_claims": [],
            "literature_evidence": [], "human_author_evidence": [],
            "support_category": "project_evidence_only", "analysis_type": "AUTHOR_ANALYSIS",
            "confidence": "medium", "planned_sections": ["3"],
            "reasoning": "源文、历史初译、历史终译与用户重译动作均有记录，但 finding、repair rationale 与作者说明缺失。",
            "counterargument": "终译仍为英文题名，因此该案例证明的是对应关系修复，不足以单独证明翻译成品质量。",
        },
        {
            "claim_id": "C2",
            "claim": "0272 将‘五个字’改为‘这句话’，把源文显式的数字回指转换为对整句言语的指称，避免‘五个字’与七字中文引语在目标语表层不一致。",
            "research_question": "RQ1", "project_evidence": [AUTHENTIC[1]],
            "synthetic_contrast_evidence": [], "literature_claims": [],
            "literature_evidence": [], "human_author_evidence": [],
            "support_category": "project_evidence_only", "analysis_type": "AUTHOR_ANALYSIS",
            "confidence": "medium", "planned_sections": ["3"],
            "reasoning": "实际差异只涉及‘五个字’至‘这句话’，修订效果可由源文与两版译文直接比较；修订动机仍未知。",
            "counterargument": "没有同期 finding 或作者解释，不能把文本效果写成已记录的主观意图。",
        },
        {
            "claim_id": "C3",
            "claim": "SC-0141 的受控对比显示，中性疑问结构与口语反问结构可以在命题内容近似不变时改变语用力度，但该对比不代表作者历史初译。",
            "research_question": "RQ2", "project_evidence": [],
            "synthetic_contrast_evidence": [SYNTHETIC], "literature_claims": [],
            "literature_evidence": [], "human_author_evidence": [],
            "support_category": "author_analysis", "analysis_type": "AUTHOR_ANALYSIS",
            "confidence": "medium", "planned_sections": ["3"],
            "reasoning": "模拟初译合理性、错误实质性与优化有效性均通过既有 synthetic gate。",
            "counterargument": "案例可能被读成模型制造并修复自身错误，因此只能作为补充性受控对比。",
        },
        {
            "claim_id": "C4",
            "claim": "三项材料形成从段落对应、话语指称到语用力度的递进，但历史证据与合成证据支持的结论强度必须分开。",
            "research_question": "RQ2", "project_evidence": AUTHENTIC,
            "synthetic_contrast_evidence": [SYNTHETIC], "literature_claims": [],
            "literature_evidence": [], "human_author_evidence": [],
            "support_category": "project_evidence_only", "analysis_type": "AUTHOR_ANALYSIS",
            "confidence": "medium", "planned_sections": ["3"],
            "reasoning": "真实案例支持实际修订事实，合成案例只支持一种经验证的可能失败机制。",
            "counterargument": "若两类证据并列为同等案例，章节会产生方法突变感。",
        },
    ]
    return _stamp({"schema_version": academic_writer.VERSIONS["argument_plan_version"],
                   "claims": claims, "planner_fallback": False,
                   "rejected_source_only_support": 0})


def _outline() -> dict[str, Any]:
    section = {
        "section_id": "3", "title": "翻译问题的层级化修订与受控对比",
        "purpose": (
            "以源译对应、话语指称和语用力度为递进主线，先分析两项真实修订，"
            "再以 SC-0141 作为补充性受控对比；须在本章内同时披露合成方法和局限。"),
        "research_questions": ["RQ1", "RQ2"],
        "claims": ["C1", "C2", "C3", "C4"],
        "cases": [*AUTHENTIC, SYNTHETIC],
        "case_groups": {"authentic_revision": AUTHENTIC,
                        "synthetic_contrast": [SYNTHETIC]},
        "literature_claims": [], "literature_evidence": [],
        "literature_sources": [], "required_statistics": [],
        "target_words": 2600, "minimum_chars": 1800,
        "allowed_conclusions": [
            "0209 只证明源译对应关系修复，不得将其归类为指称清晰度；不证明题名保留策略或最终翻译质量",
            "0272 只证明实际词语替换及其可观察的指称效果，不声称历史动机；源文写 Those five words，但不得声称可见英文引语恰好五词；中文引语为七个汉字",
            "SC-0141 只展示一种合理失败模式，不证明人类错误频率",
            "没有 Literature Evidence 时不得提及具体理论名称",
        ],
        "mandatory_structure": [
            "3.1 证据边界与分析路径", "3.2 真实修订案例：对应关系与话语指称",
            "3.3 合成对比案例：语用力度的补充性受控检验", "3.4 跨案例综合与局限",
        ],
    }
    return _stamp({
        "schema_version": academic_writer.VERSIONS["outline_version"],
        "sections": [section], "planner_fallback": False,
        "case_count_policy": {"status": "two_case_fallback", "preferred": 3,
                              "minimum": 2, "selected": 3,
                              "scarcity_disclosure": "现有项目证据仅支持两个合格真实修订案例；未用弱证据补足第三个历史案例。"},
        "case_groups": {"authentic_revision": AUTHENTIC,
                        "synthetic_contrast": [SYNTHETIC]},
    })


def _selected(canonical: dict[str, Any]) -> dict[str, Any]:
    selected = copy.deepcopy(canonical)
    mapping = {
        AUTHENTIC[0]: (["C1", "C4"], ["RQ1", "RQ2"], "core_authentic"),
        AUTHENTIC[1]: (["C2", "C4"], ["RQ1", "RQ2"], "core_authentic"),
        SYNTHETIC: (["C3", "C4"], ["RQ2"], "supplemental_controlled_contrast"),
    }
    for case in selected["cases"]:
        case["supports_claims"], case["research_questions"], case["composition_role"] = \
            mapping[case["case_id"]]
    selected["recommended_composition"] = "configuration_c"
    return _stamp(selected)


def _composition_plan(research: dict[str, Any], arguments: dict[str, Any],
                      outline: dict[str, Any], plans: dict[str, Any]) -> dict[str, Any]:
    return _stamp({
        "schema_version": PILOT_VERSION,
        "chapter_argument": (
            "翻译修订首先依赖可靠的源译对应，其次处理跨语言的话语指称；"
            "当历史修订不足以覆盖语用层问题时，明确标注的受控合成对比可作为补充，"
            "但不能与历史证据等量齐观。"),
        "case_contributions": {
            "0209": "显示源译段错配的实际纠正，并暴露最终译文质量仍不可判定的边界。",
            "0272": "显示源语词数回指转为目标语整句指称的实际修订。",
            "SC-0141": "补充真实案例未覆盖的疑问句语用力度对比。",
        },
        "configuration_intent": "C: 两项真实修订为核心，SC-0141 置于补充性方法演示小节。",
        "human_evidence_boundary": "awaiting_author_input; no answers used",
        "literature_evidence_boundary": "no grounded literature evidence; no theory names allowed",
        "research_model": research, "argument_plan": arguments,
        "outline": outline, "case_analysis_plans": plans,
    })


def _ensure_chapter_frame(content: str) -> str:
    content = content.strip()
    while True:
        trimmed = re.sub(
            r"\A(?:#\s+第三章[^\n]*\n+|"
            r"##\s+3[.．、]?\s+[^\n]*\n+)", "", content)
        if trimmed == content:
            break
        content = trimmed
    content = re.sub(r"^###\s+(3\.\d+\.\d+\s+.*)$", r"#### \1",
                     content, flags=re.MULTILINE)
    content = re.sub(r"^##\s+(3\.\d+\s+.*)$", r"### \1",
                     content, flags=re.MULTILINE)
    return (
        "# 第三章 翻译问题的层级化修订与受控对比\n\n"
        "## 3 翻译问题的层级化修订与受控对比\n\n" + content + "\n")


def _review(api_key: str, provider: str, model: str, report: str,
            validation: dict[str, Any], quality: dict[str, Any],
            plan: dict[str, Any]) -> dict[str, Any]:
    system = (
        "你是独立的 MTI 导师式审稿人。只审查提供的第三章成稿和结构化证据，不改写正文，"
        "不补造文献、译者意图或 Human Evidence。0209/0272 是历史修订；SC-0141 是"
        "分析阶段生成的受控合成对比。请检查正文实际句子并只输出 JSON。"
        "dimensions 必须逐项给出 status(pass/pass_with_warnings/review_required/fail/"
        "not_applicable) 与 reason：argument_coherence, case_analytical_depth, "
        "case_complementarity, theory_case_fit, evidence_utilization, "
        "synthetic_methodology_clarity, cross_case_progression, academic_specificity, "
        "redundancy, conclusion_discipline, provenance_correctness。"
        "另输出 supervisor_review、sc0141_adversarial_review、configuration_comparison(A/B/C)、"
        "recommended_configuration、passage_inspection、readiness。passage_inspection 必须含"
        "strongest_analytical_passage、weakest_analytical_passage、best_authentic_reasoning_chain、"
        "best_synthetic_reasoning_chain、most_awkward_transition、most_generic_paragraph；"
        "每项给出正文中的逐字短摘录 excerpt 和 observation。不得给单一总分。"
        "SC-0141 的 adversarial review 必须回答合理性、合格译者是否可能受诱、错误实质性、"
        "优化是否因诊断理由而更好、是在展示翻译现象还是模型自导自演、学术增量、"
        "contribution(unique/partially_redundant/redundant/methodologically_disruptive)、"
        "retain 与 role(core/supplement/remove)。配置比较分别评估 coherence/depth/"
        "methodological_clarity/diversity/provenance_risk/reader_burden/naturalness。"
        "三种配置的定义不得改写：A=0209+0272（仅两项真实修订）；"
        "B=0209+0272+SC-0141，且 SC-0141 作为并列核心案例；"
        "C=0209+0272 作为核心，SC-0141 放在补充性/方法演示小节。"
        "readiness 只能取 engineering_valid/academically_reviewable/"
        "supervisor_review_required/submission_candidate。"
    )
    raw = _json_call(api_key, system, {
        "chapter": report,
        "composition_plan": {k: plan[k] for k in (
            "chapter_argument", "case_contributions", "configuration_intent",
            "human_evidence_boundary", "literature_evidence_boundary")},
        "deterministic_validation": validation,
        "existing_quality_review": quality,
    }, provider, model)
    return _stamp({"schema_version": REVIEW_VERSION, **raw})


def _supervisor_brief(review: dict[str, Any], plan: dict[str, Any]) -> str:
    supervisor = review.get("supervisor_review") or {}
    if not isinstance(supervisor, dict):
        supervisor = {"overall_assessment": str(supervisor)}
    adversarial = review.get("sc0141_adversarial_review") or {}
    if not isinstance(adversarial, dict):
        adversarial = {"academic_increment": str(adversarial)}

    def bullets(value: Any) -> str:
        items = value if isinstance(value, list) else [value] if value else []
        return "\n".join(f"- {item}" for item in items) or "- 无。"

    inspection = review.get("passage_inspection") or {}
    dimensions = review.get("dimensions") or {}
    strongest = [
        (inspection.get("best_authentic_reasoning_chain") or {}).get("observation"),
        (inspection.get("best_synthetic_reasoning_chain") or {}).get("observation"),
    ]
    weak_points = [
        value.get("reason") for value in dimensions.values()
        if isinstance(value, dict) and value.get("status") in {
            "pass_with_warnings", "review_required", "fail"}
    ]
    why_synthetic = (
        f"保留为补充性受控对比（{adversarial.get('contribution', '未分类')}）："
        f"{adversarial.get('academic_increment', '')}")

    return f"""# Chapter 3 Supervisor Review Brief

## Chapter 3 argument

{plan['chapter_argument']}

## Why 0209 was selected

{plan['case_contributions']['0209']}

## Why 0272 was selected

{plan['case_contributions']['0272']}

## Why SC-0141 was or was not retained

{why_synthetic}

## Authentic vs synthetic methodology

两项真实修订作为核心证据；SC-0141 仅作补充性受控对比，不与历史修订等量齐观。

## Strongest analytical findings

{bullets(supervisor.get('strongest_findings') or strongest)}

## Remaining weak points

{bullets(supervisor.get('remaining_weak_points') or weak_points)}

## Theory questions

{bullets(supervisor.get('theory_questions') or ['当前没有 Literature Evidence；是否在后续版本登记并核验可真正解释话语指称或语用力度的理论来源？'])}

## Human Evidence still needed

{bullets(supervisor.get('human_evidence_still_needed') or ['0209：保留英文题名及初译错配原因。', '0272：“五个字”改为“这句话”的实际修订考虑。'])}

## Questions requiring supervisor judgment

{bullets(supervisor.get('questions_requiring_supervisor_judgment') or ['是否接受配置 C：两项真实修订为核心，SC-0141 仅作补充性受控对比？', '在获得 Human Evidence 前，0209 是否只保留为证据边界案例？'])}
"""


def _set_font(run, *, latin="Songti SC", east_asia="Songti SC", size=11,
              color="000000", bold=False, italic=False) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    _set_font(run, size=9, color="666666")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    tail = paragraph.add_run(" 页")
    _set_font(tail, size=9, color="666666")


def _docx(md: str, path: Path, model: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = Inches(1)
    section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Songti SC"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for name, size, color, before, after in (
            ("Heading 1", 16, "2E74B5", 18, 10),
            ("Heading 2", 13, "2E74B5", 12, 6),
            ("Heading 3", 12, "1F4D78", 8, 4)):
        style = doc.styles[name]
        style.font.name = "Heiti SC"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Heiti SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_font(header.add_run("MTI Translation Practice Report | Chapter 3 Pilot"),
              size=9, color="666666")
    _page_field(section.footer.paragraphs[0])

    visible = re.sub(r"<!--.*?-->", "", md, flags=re.DOTALL)
    blocks = re.split(r"\n\s*\n", visible.strip())
    quote_labels = {"SOURCE": "源文", "INITIAL": "历史初译", "TARGET": "历史终译",
                    "SYNTHETIC_SOURCE": "真实源文", "SIMULATED": "模拟初译",
                    "OPTIMIZED": "优化译文"}
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            _set_font(p.add_run(block[2:]), latin="Heiti SC", east_asia="Heiti SC", size=22,
                      color="203748", bold=True)
            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.paragraph_format.space_after = Pt(18)
            _set_font(meta.add_run(f"隔离式写作试运行 | OpenCode Go · {model}"),
                      size=10, color="666666", italic=True)
            continue
        if block.startswith("#### "):
            doc.add_heading(block[5:], level=3)
            continue
        if block.startswith("### "):
            doc.add_heading(block[4:], level=2)
            continue
        if block.startswith("## "):
            doc.add_heading(block[3:], level=1)
            continue
        quote_lines = [line.strip() for line in block.splitlines() if line.strip()]
        quote_matches = [re.fullmatch(
            r">\s*\[(SOURCE|INITIAL|TARGET|SYNTHETIC_SOURCE|SIMULATED|OPTIMIZED)\s+"
            r"([^\]]+)\]:\s*(.*)", line) for line in quote_lines]
        if quote_lines and all(quote_matches):
            for match in quote_matches:
                kind, case_id, text = match.groups()
                short_id = case_id.rsplit("-", 1)[-1] \
                    if kind in {"SOURCE", "INITIAL", "TARGET"} else case_id
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.28)
                p.paragraph_format.right_indent = Inches(0.18)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.208
                _set_font(p.add_run(f"{quote_labels[kind]}（{short_id}）："),
                          size=10.5, color="1F4D78", bold=True)
                _set_font(p.add_run(text), size=10.5, color="333333")
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        friendly = re.sub(r"seg-[A-Za-z0-9_-]+-(\d{4})", r"\1", block)
        _set_font(p.add_run(re.sub(r"\*\*(.*?)\*\*", r"\1", friendly.replace("\n", " "))),
                  size=11)
    doc.core_properties.title = "第三章 翻译问题的层级化修订与受控对比"
    doc.core_properties.subject = "MTI Chapter 3 composition pilot"
    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--job-id", default=JOB)
    parser.add_argument("--out-dir", required=True)
    parser.add_argument("--synthetic-dir", required=True)
    parser.add_argument("--human-pilot-dir", required=True)
    parser.add_argument("--provider", default="OpenCode Go")
    parser.add_argument("--model", default="glm-5.2")
    parser.add_argument(
        "--skip-docx", action="store_true",
        help="omit DOCX when the current renderer cannot pass visual QA")
    args = parser.parse_args()
    if args.job_id != JOB:
        parser.error(f"this frozen pilot only supports {JOB}")
    api_key = os.environ.get("OPENCODE_GO_API_KEY", "")
    if not api_key:
        parser.error("OPENCODE_GO_API_KEY is required")
    out_dir = Path(args.out_dir)
    if out_dir.exists():
        parser.error("out-dir already exists")
    state_path = Path("outputs") / args.job_id / "state.json"
    synthetic_dir = Path(args.synthetic_dir)
    human_dir = Path(args.human_pilot_dir)
    canonical_selected = _load(synthetic_dir / "selected-cases.json")
    human_state = _load(human_dir / "human-evidence-pilot-state.json")
    if [x.get("case_id") for x in canonical_selected.get("cases", [])] != [*AUTHENTIC, SYNTHETIC]:
        parser.error("canonical case selection does not match the frozen pilot")
    if human_state.get("status") != "awaiting_author_input" or human_state.get("phase_b_started"):
        parser.error("Human Evidence boundary is not awaiting_author_input")
    synthetic_artifact = academic_writer._read_artifact(
        synthetic_dir / "synthetic-case-validation.jsonl")
    synthetic_case = next((x for x in (synthetic_artifact or {}).get("items", [])
                           if x.get("case_id") == SYNTHETIC), None)
    if not synthetic_case or not synthetic_case.get("validation", {}).get("academic_case_eligible"):
        parser.error("SC-0141 is absent or no longer eligible")

    out_dir.mkdir(parents=True)
    before_hash = _sha256(state_path)
    state = _load(state_path)
    evidence = academic_evidence.build_academic_evidence(state, args.job_id)
    segments = academic_evidence.segment_index(evidence)
    if any(case_id not in segments or not academic_evidence.is_eligible_revision_case(
            segments[case_id])
           for case_id in AUTHENTIC):
        raise RuntimeError("an authentic case failed the existing revision gate")
    if synthetic_case["source_text"] not in segments[
            synthetic_case["source_segment_id"]]["source"]:
        raise RuntimeError("SC-0141 no longer matches the real source segment")

    research = _research_model()
    arguments = _argument_plan()
    selected = _selected(canonical_selected)
    outline = _outline()
    empty_sources = {"schema_version": academic_writer.VERSIONS["literature_sources_version"],
                     "sources": []}
    empty_evidence = {"schema_version": academic_writer.VERSIONS["literature_evidence_version"],
                      "items": []}
    empty_claims = {"schema_version": academic_writer.VERSIONS["literature_claims_version"],
                    "items": []}
    plans = case_analysis.build_case_analysis_plans(
        evidence, selected, arguments, empty_claims, core.call_llm,
        args.provider, api_key, args.model, human_evidence=[])
    plan_by_id = {x["case_id"]: x for x in plans["plans"]}
    plan_by_id[AUTHENTIC[0]].update({
        "problem": {"type": "information_structure",
                    "statement": "历史初译文本与本段源文在语言和内容上都无对应关系。",
                    "grounded": True},
        "decision_rationale": (
            "可观察到的修订是将无对应关系的中文片段整体替换为源文题名，"
            "客观上恢复了当前段的源译对应；但保留英文题名的原因未记录。"),
        "translation_effect": {
            "dimension": "information_structure",
            "demonstrated_by": (
                "‘那是夏天，或是夏末。那是午后。’与题名源文无对应，"
                "终译‘RIOT IN CELL BLOCK 11’与本段源文一致。")},
        "bounded_conclusion": (
            "该修订恢复了本段的源译对应，但现有证据既不说明初译错配的"
            "生成机制，也不证明保留英文题名是有意翻译策略。"),
        "recommended_human_evidence": [
            "保留英文题名而不译出的原因。",
            "初译为何出现与本段无对应关系文本的工作流记录。"],
    })
    plan_by_id[AUTHENTIC[1]]["decision_rationale"] = (
        "可观察到的‘五个字’至‘这句话’替换，客观上消除了‘五个字’与"
        "七字中文引语的表层不一致，使回指对象变为整句言语；历史动机未记录。")
    plan_by_id[AUTHENTIC[1]].update({
        "problem": {
            "type": "reference_resolution",
            "statement": (
                "源文显式使用 ‘Those five words’ 回指前文引语，但不得据此声称"
                "可见英文引语恰好五词；在汉译中，‘你不会经历战争’为七个汉字，"
                "初译‘这五个字’因此在目标语表层产生可直接核对的不一致。"),
            "grounded": True,
        },
        "translation_effect": {
            "dimension": "reference_clarity",
            "demonstrated_by": (
                "将‘五个字’替换为‘这句话’，使回指不再依赖与七字中文引语不符的"
                "目标语计数，而改为指向整句言语。"),
        },
        "bounded_conclusion": (
            "本案例只证明‘五个字’至‘这句话’的实际替换及其消除目标语表层"
            "计数不一致的效果；不解释源文为何使用 five，也不推断修订者动机。"),
        "alternatives": [
            {
                "label": "analytical_comparison",
                "text": (
                    "若使用‘这番话’，同样可把回指对象从字数转为整句言语，"
                    "但语体更书面；历史记录中的实际终译并非此方案。"),
            },
            {
                "label": "counterfactual_rendering",
                "text": (
                    "若保留‘这五个字’，目标语表层仍与七字中文引语不一致；"
                    "历史记录未见连带改写引语的方案。"),
            },
        ],
    })
    plans = _stamp({k: v for k, v in plans.items() if k != "content_hash"})
    composition = _composition_plan(research, arguments, outline, plans)
    _write_json(out_dir / "chapter-3-composition-plan.json", composition)

    section_plan = outline["sections"][0]
    packet = academic_writer._section_packet(
        section_plan, research, arguments, selected, evidence, outline, [],
        empty_sources, empty_evidence, empty_claims, plans)
    content = academic_writer._write_section(
        packet, core.call_llm, args.provider, api_key, args.model)
    report = _ensure_chapter_frame(academic_writer.finalize_report_tokens(
        academic_writer._compose_report([{
            "section_id": "3", "title": section_plan["title"], "content": content,
        }]), evidence, selected))
    validation = academic_validator.validate_academic_report(
        report, evidence, research, arguments, selected, outline,
        empty_sources, empty_evidence, empty_claims, human_evidence=[],
        synthetic_artifact=synthetic_artifact)
    if validation.get("status") == "fail":
        repair_issues = [dict(item, section_id="3") for item in validation.get("issues", [])]
        content = academic_writer._write_section(
            packet, core.call_llm, args.provider, api_key, args.model,
            repair_issues=repair_issues, existing=content)
        report = _ensure_chapter_frame(academic_writer.finalize_report_tokens(
            academic_writer._compose_report([{
                "section_id": "3", "title": section_plan["title"], "content": content,
            }]), evidence, selected))
        validation = academic_validator.validate_academic_report(
            report, evidence, research, arguments, selected, outline,
            empty_sources, empty_evidence, empty_claims, human_evidence=[],
            synthetic_artifact=synthetic_artifact)
    sections = [{"section_id": "3", "title": section_plan["title"],
                 "content": content}]
    quality = academic_quality.evaluate_quality(
        research, arguments, selected, outline, sections, evidence,
        empty_sources, empty_evidence, empty_claims, validation,
        core.call_llm, args.provider, api_key, args.model,
        case_analysis_plans=plans)
    repairable_quality = [
        item for item in quality.get("findings", [])
        if item.get("priority") in {"P0", "P1", "P2"}
        and item.get("section_id") == "3"]
    if _has_invalid_0272_count_claim(content):
        repairable_quality.append({
            "issue_id": "PILOT-FACT-001", "section_id": "3",
            "claim_id": "C2", "case_id": AUTHENTIC[1], "priority": "P1",
            "severity": "high", "repair_action": "narrow_claim",
            "reason": (
                "可见英文引语 You will not have a war 按空格分词并非五词。"
                "只能陈述源文后句显式写 Those five words，以及中文引语"
                "为七个汉字；不得声称源文引语恰好五词，也不得建议"
                "用注释说明‘源文为五个英文词’。"),
            "recommended_action": (
                "删除不可验证的五词断言及相关反事实备选；仅分析"
                "‘五个字’与七字中文引语的目标语表层不一致。"),
        })
    if repairable_quality:
        content = academic_writer._write_section(
            packet, core.call_llm, args.provider, api_key, args.model,
            repair_issues=repairable_quality, existing=content)
        report = _ensure_chapter_frame(academic_writer.finalize_report_tokens(
            academic_writer._compose_report([{
                "section_id": "3", "title": section_plan["title"], "content": content,
            }]), evidence, selected))
        validation = academic_validator.validate_academic_report(
            report, evidence, research, arguments, selected, outline,
            empty_sources, empty_evidence, empty_claims, human_evidence=[],
            synthetic_artifact=synthetic_artifact)
        sections = [{"section_id": "3", "title": section_plan["title"],
                     "content": content}]
        quality = academic_quality.evaluate_quality(
            research, arguments, selected, outline, sections, evidence,
            empty_sources, empty_evidence, empty_claims, validation,
            core.call_llm, args.provider, api_key, args.model,
            case_analysis_plans=plans)
    if _has_invalid_0272_count_claim(content):
        raise RuntimeError("0272 output still contains an invalid word/character count claim")
    (out_dir / "chapter-3-final-pilot.md").write_text(report, encoding="utf-8")
    _write_json(out_dir / "chapter-3-validation.json", validation)
    semantic_review = academic_writer._semantic_review(
        report, research, arguments, outline, selected,
        core.call_llm, args.provider, api_key, args.model)
    special_review = _review(api_key, args.provider, args.model, report,
                             validation, quality, composition)
    combined_quality = _stamp({
        "schema_version": "chapter3-quality-review-v1",
        "existing_academic_quality": quality,
        "existing_semantic_review": semantic_review,
        "chapter_composition_review": special_review,
    })
    _write_json(out_dir / "chapter-3-quality-review.json", combined_quality)
    (out_dir / "chapter-3-supervisor-review-brief.md").write_text(
        _supervisor_brief(special_review, composition), encoding="utf-8")

    after_hash = _sha256(state_path)
    provenance = _stamp({
        "schema_version": "chapter3-provenance-summary-v1",
        "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "job_id": args.job_id,
        "baseline_commit": subprocess.run(
            ["git", "rev-parse", "HEAD"], check=True, text=True,
            capture_output=True).stdout.strip(),
        "historical_state": {"path": str(state_path), "sha256_before": before_hash,
                             "sha256_after": after_hash,
                             "unchanged": before_hash == after_hash},
        "canonical_sources": {
            "authentic_revision_audit": str(human_dir / "revision-case-audit.json"),
            "human_evidence_state": str(human_dir / "human-evidence-pilot-state.json"),
            "synthetic_selection": str(synthetic_dir / "selected-cases.json"),
            "synthetic_validation": str(synthetic_dir / "synthetic-case-validation.jsonl"),
        },
        "case_verification": {
            "0209": {"case_type": "authentic_revision", "eligibility": "eligible",
                     "finding": "not_recorded", "repair_history": "not_recorded",
                     "human_action": "retranslated", "human_rationale": "awaiting_author_input",
                     "analytical_boundary": "alignment repair only; title retention intent unknown"},
            "0272": {"case_type": "authentic_revision", "eligibility": "eligible",
                     "finding": "not_recorded", "repair_history": "not_recorded",
                     "human_action": "not_recorded", "human_rationale": "awaiting_author_input",
                     "analytical_boundary": "observable lexical and referential effect only"},
            "SC-0141": {"case_type": "synthetic_contrast", "eligibility": "eligible",
                        "historical": False, "generated_for_analysis": True,
                        "baseline_plausibility": synthetic_case["baseline_plausibility"]["status"],
                        "repair_correctness": synthetic_case["validation"]["repair_correctness"]},
        },
        "evidence_boundaries": {"human_evidence_status": "awaiting_author_input",
                                "human_evidence_entries_used": 0,
                                "literature_sources_used": 0,
                                "theory_connections": "not_applicable"},
        "provider": args.provider, "model": args.model,
        "endpoint_class": "OpenAI-compatible chat completions",
        "versions": {"prompt": PILOT_VERSION,
                     "writer": academic_writer.VERSIONS["writer_version"],
                     "reviewer": REVIEW_VERSION,
                     "semantic_reviewer": academic_writer.VERSIONS["reviewer_version"],
                     "quality_reviewer": academic_writer.VERSIONS["academic_quality_version"],
                     "validator": academic_validator.VALIDATOR_VERSION,
                     "case_analysis": case_analysis.ANALYSIS_VERSION},
        "recommended_configuration": special_review.get("recommended_configuration"),
        "readiness": special_review.get("readiness"),
        "docx_delivery": ({
            "status": "omitted_unsupported_current_renderer",
            "reason": "visual QA could not render CJK glyphs in the current LibreOffice environment",
        } if args.skip_docx else {
            "status": "generated_pending_visual_qa",
            "preset": "narrative_proposal",
            "header_pattern": "editorial_cover_restrained",
            "named_overrides": {"east_asia_body_font": "Songti SC",
                                "east_asia_heading_font": "Heiti SC",
                                "title": "22pt #203748 centered",
                                "quote": "10.5pt indented"},
        }),
    })
    _write_json(out_dir / "chapter-3-provenance-summary.json", provenance)
    if not args.skip_docx:
        _docx(report, out_dir / "chapter-3-final-pilot.docx", args.model)
    run = _stamp({
        "schema_version": "chapter3-pilot-run-v1", "status": "complete"
        if validation.get("status") != "fail" and before_hash == after_hash else "failed",
        "provider": args.provider, "model": args.model,
        "validation_status": validation.get("status"),
        "quality_readiness": special_review.get("readiness"),
        "historical_state_unchanged": before_hash == after_hash,
        "artifact_files": sorted(path.name for path in out_dir.iterdir()),
    })
    _write_json(out_dir / "run-manifest.json", run)
    print(out_dir)
    return 0 if run["status"] == "complete" else 3


if __name__ == "__main__":
    raise SystemExit(main())
