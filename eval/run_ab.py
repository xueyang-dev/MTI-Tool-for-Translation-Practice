"""受控 A/B/C/D 评测编排器（TransPraxis / 译践 Evaluation Harness 入口）。

实验矩阵（Governance × Reviewed TM）：
    A = baseline                            + 无 reviewed TM
    B = current + quality mode（治理栈）   + 无 reviewed TM
    C = pre-governance                    + 有 reviewed TM
    D = current + quality mode            + 有 reviewed TM

回答：A→B 治理增量；A→C TM 增量；B→D 治理上再加 TM；A→D 完整工作流。

数据卫生：
- 真实语料（source.bin / 译文 / TM / 抽样正文）只写入 eval/results/（gitignored），
  一律不入库；入库的只有工具代码、聚合指标与合成 fixture。
- A 臂与其他运行臂使用同一份当前代码；差异由运行配置和 TM policy 控制。

用法：
    python eval/run_ab.py --config eval/config.example.json [--mock]
    python eval/run_ab.py --config ... --segments 0:500
    python eval/run_ab.py --config ... --glossary-from-job <local-job-id> --lock
"""
from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List

ROOT = Path(__file__).resolve().parent.parent
EVAL_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

import core  # noqa: E402   （corpus 构建与指标依赖当前代码）

from eval import config as eval_config  # noqa: E402
from eval.metrics import aggregate, qa_density, terminology, workflow  # noqa: E402
from eval.reporting import csv_report, json_report, markdown_report  # noqa: E402
from eval.sampling import blind_review  # noqa: E402


# ---------------- 语料 ----------------

def _docx_paragraphs(file_bytes: bytes) -> List[str]:
    import io
    import re
    from docx import Document
    paras: List[str] = []
    doc = Document(io.BytesIO(file_bytes))
    for p in doc.paragraphs:
        for sub in re.split(r"\n+", p.text):
            t = sub.strip()
            if len(t) > 1:
                paras.append(t)
    return paras


def build_corpus(job_id: str, subset: List[int], out_dir: Path) -> Dict[str, Any]:
    """从 outputs/<job>/source.bin 构建确定性子集（corpus.docx + corpus.jsonl）。

    corpus.jsonl 含原文文本：只写入 eval/results/（gitignored），禁止入库。
    """
    src = ROOT / "outputs" / job_id / "source.bin"
    if not src.is_file():
        raise FileNotFoundError(f"找不到语料源文件：{src}")
    file_bytes = src.read_bytes()
    name = (ROOT / "outputs" / job_id / "state.json").read_text(encoding="utf-8")
    filename = json.loads(name).get("filename", "?")
    if filename.lower().endswith(".pdf"):
        paragraphs = [core.clean_xml_chars(p)
                      for p in core.extract_pdf_paragraphs(file_bytes)]
    elif filename.lower().endswith(".docx"):
        paragraphs = _docx_paragraphs(file_bytes)
    else:
        raise ValueError(f"不支持的语料类型：{filename}")
    if not paragraphs:
        raise ValueError("语料无文本层（扫描版 PDF 需先 OCR）")
    start, end = subset
    if start == 0 and end == 0:
        selected = paragraphs
    else:
        selected = paragraphs[start:end]
    if not selected:
        raise ValueError(f"子集为空：[{start}, {end})，全文 {len(paragraphs)} 段")

    out_dir.mkdir(parents=True, exist_ok=True)
    from docx import Document
    doc = Document()
    for p in selected:
        doc.add_paragraph(p)
    docx_path = out_dir / "corpus.docx"
    doc.save(docx_path)
    meta = {
        "job_id": job_id,
        "filename": filename,
        "subset": subset,
        "full_segments": len(paragraphs),
        "selected_segments": len(selected),
        "source_chars": sum(len(p) for p in selected),
    }
    with (out_dir / "corpus.jsonl").open("w", encoding="utf-8") as f:
        for i, p in enumerate(selected):
            f.write(json.dumps({"index": i, "source": p}, ensure_ascii=False) + "\n")
    (out_dir / "corpus_meta.json").write_text(
        json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return meta


def load_corpus_segments(out_dir: Path) -> List[str]:
    segs: List[str] = []
    with (out_dir / "corpus.jsonl").open(encoding="utf-8") as f:
        for line in f:
            segs.append(json.loads(line)["source"])
    return segs


# ---------------- 基线 worktree ----------------

def ensure_baseline_worktree(ref: str) -> Path:
    dest = EVAL_DIR / ".worktrees" / "baseline"
    if (dest / ".git").is_file() or (dest / ".git").is_dir():
        cur = subprocess.run(["git", "rev-parse", "HEAD"], cwd=dest,
                             capture_output=True, text=True)
        if cur.returncode == 0 and cur.stdout.strip() == ref:
            return dest
    subprocess.run(["git", "worktree", "remove", str(dest), "--force"],
                   capture_output=True, text=True)
    subprocess.run(["git", "worktree", "add", "--detach", str(dest), ref],
                   check=True, capture_output=True, text=True)
    return dest


# ---------------- 单臂运行 ----------------

def run_arm(arm: str, *, code_root: Path, run_dir: Path, corpus_docx: Path,
            glossary_file: Path, tm_seed: Path | None, mock: bool,
            run_cfg: Dict[str, Any], api_key: str) -> Dict[str, Any]:
    run_dir.mkdir(parents=True, exist_ok=True)
    cmd = [sys.executable, str(EVAL_DIR / "runner.py"),
           "--run-dir", str(run_dir),
           "--corpus", str(corpus_docx),
           "--glossary", str(glossary_file),
           "--arm", arm,
           "--code-ref", _git_head(code_root),
           "--provider", run_cfg["provider"],
           "--model", run_cfg["model"],
           "--target-lang", run_cfg["target_lang"],
           "--theory", run_cfg["translation_theory"],
           "--style-rules", run_cfg["style_rules"]]
    if run_cfg.get("enable_review"):
        cmd.append("--enable-review")
    if mock:
        cmd.append("--mock")
    if arm in ("C", "D") and tm_seed is not None:
        cmd += ["--tm-seed", str(tm_seed)]
    env = dict(os.environ)
    env["PYTHONPATH"] = str(code_root)
    if api_key:
        env["TRANSPRAXIS_EVAL_API_KEY"] = api_key
    proc = subprocess.run(cmd, cwd=str(code_root), env=env,
                          capture_output=True, text=True, timeout=8 * 3600)
    if proc.returncode != 0:
        tail = (proc.stderr or proc.stdout)[-2000:]
        raise RuntimeError(f"arm {arm} 运行失败（code={proc.returncode}）：\n{tail}")
    meta_path = run_dir / "run_meta.json"
    run_meta = json.loads(meta_path.read_text(encoding="utf-8")) \
        if meta_path.is_file() else {}
    return run_meta


def _git_head(code_root: Path) -> str:
    proc = subprocess.run(["git", "rev-parse", "HEAD"], cwd=code_root,
                          capture_output=True, text=True)
    return proc.stdout.strip() if proc.returncode == 0 else "unknown"


# ---------------- 历史术语表导出（local-only） ----------------

def export_glossary_from_job(job_id: str, lock: bool, out_path: Path) -> Path:
    sp = ROOT / "outputs" / job_id / "state.json"
    if not sp.is_file():
        raise FileNotFoundError(f"找不到任务 state：{sp}")
    state = json.loads(sp.read_text(encoding="utf-8"))
    auto_terms = state.get("auto_terms") or {}
    entries = [{
        "source": k,
        "target": v,
        "preferred": v,
        "behavior": "translate",
        "status": "locked" if lock else "candidate",
        "scope": "global",
        "note": ("eval 自动锁定（未经人工审核）" if lock else "eval 候选"),
    } for k, v in sorted(auto_terms.items()) if k and v]
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps({
        "source_job": job_id,
        "human_reviewed": False,
        "auto_locked": bool(lock),
        "entries": entries,
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    return out_path


# ---------------- 共享收尾（指标 / 增量 / 盲评包 / 报告）----------------

def load_existing_runs(out_dir: Path):
    """从已有 out_dir/runs/<arm>/<job>/state.json 恢复运行结果（report-only）。"""
    states: Dict[str, Dict[str, Any]] = {}
    metas: Dict[str, Dict[str, Any]] = {}
    for arm in ("A", "B", "C", "D"):
        meta_p = out_dir / "runs" / arm / "run_meta.json"
        if not meta_p.is_file():
            continue
        meta = json.loads(meta_p.read_text(encoding="utf-8"))
        job = meta.get("job_id")
        sp = (out_dir / "runs" / arm / job / "state.json") if job else None
        if sp is not None and sp.is_file():
            states[arm] = json.loads(sp.read_text(encoding="utf-8"))
            metas[arm] = meta
    return states, metas


def _metrics_for_states(states: Dict[str, Dict[str, Any]],
                        glossary_entries: List[Dict[str, Any]]) -> Dict[str, Any]:
    runs: Dict[str, Any] = {}
    for arm, state in states.items():
        runs[arm] = {
            "terminology": terminology.compute_terminology_metrics(
                state, glossary_entries),
            "qa": qa_density.compute_qa_density(state),
            "workflow": workflow.compute_workflow_metrics(state),
        }
    return runs


def _finalize(out_dir: Path, states: Dict[str, Dict[str, Any]],
              segments: List[str] | None,
              glossary_entries: List[Dict[str, Any]],
              meta: Dict[str, Any], seed: int) -> Dict[str, Any]:
    """共享收尾：指标 -> 增量 -> 盲评包 -> JSON/CSV/MD。"""
    runs = _metrics_for_states(states, glossary_entries)
    delta_pairs = ["A_to_B", "A_to_C", "B_to_D", "A_to_D"]
    deltas = aggregate.compute_deltas(runs, delta_pairs)
    packet_rows, key_rows = [], []
    packet_path = None
    if segments and "A" in states and "B" in states:
        packet_rows, key_rows = blind_review.sample_packet(
            states["A"], states["B"], segments, glossary_entries, seed=seed)
        packet_path, _key_path = blind_review.write_packet(
            packet_rows, key_rows, out_dir / "blind_review")
        print(f"盲评抽样包：{packet_path}（{len(packet_rows)} 段；"
              f"key 文件 local-only）")
        # 盲评包 v2：排除相同对，A/B 位置平衡
        rows2, keys2 = blind_review.sample_packet_v2(
            states["A"], states["B"], segments, glossary_entries, seed=seed)
        packet2_path, key2_path = blind_review.write_packet(
            rows2, keys2, out_dir / "blind_review", prefix="blind_review_v2")
        informative2 = sum(1 for r in rows2 if r.get("identical") == "0")
        print(f"盲评包 v2：{packet2_path}（{len(rows2)} 段，"
              f"有效差异 {informative2}，key local-only：{key2_path}）")
    report = aggregate.build_report(
        meta={**meta, "created_at":
              datetime.now(timezone.utc).isoformat(timespec="seconds")},
        runs=runs,
        deltas=deltas,
        human_review={
            "status": "pending",
            "packet": str(packet_path) if packet_path else None,
            "packet_v2": str(packet2_path) if packet2_path else None,
            "segments": len(rows2) if rows2 else len(packet_rows),
            "note": "blind review：Candidate A/B 随机映射，key 文件 local-only",
        },
    )
    json_path = json_report.write_json_report(report, out_dir)
    csv_report.write_terms_csv(runs, out_dir)
    csv_report.write_findings_csv(states, out_dir)
    md_path = markdown_report.write_markdown_report(report, out_dir)
    print(f"\n✅ evaluation-report.json -> {json_path}")
    print(f"   markdown -> {md_path}")
    return report


# ---------------- 主流程 ----------------

def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="TransPraxis / 译践 受控 A/B/C/D 评测")
    ap.add_argument("--config", default=None)
    ap.add_argument("--arms", default=None, help="如 ABCD 或 AB")
    ap.add_argument("--segments", default=None, help="如 0:300 或 all")
    ap.add_argument("--mock", action="store_true", help="离线自测（确定性 mock LLM）")
    ap.add_argument("--corpus-override", default=None,
                    help="离线自测：跳过 outputs 语料构建，使用给定 corpus.docx"
                         "（同目录需有 corpus.jsonl / corpus_meta.json）")
    ap.add_argument("--report-only", action="store_true",
                    help="不运行任何臂：从已有 out_dir/runs/* 重算指标/报告/盲评包"
                         "（用于修指标后复用已跑结果，不再消耗 API）")
    ap.add_argument("--glossary-from-job", default=None,
                    help="从 outputs/<job>/auto_terms 导出术语表（local-only）")
    ap.add_argument("--lock", action="store_true",
                    help="配合 --glossary-from-job：导出时自动锁定（明确标注未经人工审核）")
    ap.add_argument("--out", default=None)
    args = ap.parse_args(argv)

    cfg = eval_config.load_config(args.config)
    arms = [a for a in (args.arms or "".join(cfg["arms"]))]
    if args.segments:
        if args.segments == "all":
            cfg["corpus"]["subset"] = [0, 0]
        else:
            start, end = (int(x) for x in args.segments.split(":"))
            cfg["corpus"]["subset"] = [start, end]
    out_dir = Path(args.out or cfg.get("out_dir") or
                   f"eval/results/{datetime.now().strftime('%Y%m%d-%H%M%S')}").resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    print(eval_config.describe_run_matrix(arms))
    print(f"\n输出目录（local-only）：{out_dir}")

    # 术语表
    if args.glossary_from_job:
        glossary_file = (out_dir / "glossary_eval.json").resolve()
        export_glossary_from_job(args.glossary_from_job, args.lock, glossary_file)
        cfg["glossary"] = str(glossary_file)
        print(f"术语表已从历史任务导出（auto_locked={args.lock}）：{glossary_file}")
    else:
        glossary_file = Path(cfg["glossary"])
        if not glossary_file.is_file():
            raise FileNotFoundError(
                f"术语表不存在：{glossary_file}；可用 --glossary-from-job <job> [--lock] 导出")
    glossary_entries = eval_config.resolve_glossary_entries(cfg)

    if args.report_only:
        print("report-only：复用已有运行结果重算指标与报告")
        states, metas = load_existing_runs(out_dir)
        if not states:
            raise RuntimeError(f"{out_dir / 'runs'} 中没有可用的运行结果")
        segments = load_corpus_segments(out_dir) \
            if (out_dir / "corpus.jsonl").is_file() else None
        corpus_meta = {}
        cm_path = out_dir / "corpus_meta.json"
        if cm_path.is_file():
            corpus_meta = json.loads(cm_path.read_text(encoding="utf-8"))
        meta = {
            "baseline_ref": cfg["code"]["baseline_ref"],
            "current_ref": {a: m.get("code_ref") for a, m in metas.items()},
            "corpus": corpus_meta,
            "glossary": str(glossary_file),
            "glossary_auto_locked": args.lock,
            "tm_seed": cfg.get("tm_seed"),
            "run": cfg["run"],
            "mock": args.mock,
            "recomputed": True,
        }
        _finalize(out_dir, states, segments, glossary_entries, meta,
                  seed=cfg.get("seed", 42))
        return 0

    # 语料
    if args.corpus_override:
        corpus_docx = Path(args.corpus_override).resolve()
        corpus_meta = json.loads(
            (corpus_docx.parent / "corpus_meta.json").read_text(encoding="utf-8"))
        segments = load_corpus_segments(corpus_docx.parent)
    else:
        corpus_meta = build_corpus(cfg["corpus"]["job_id"],
                                   cfg["corpus"]["subset"], out_dir)
        corpus_docx = (out_dir / "corpus.docx").resolve()
        segments = load_corpus_segments(out_dir)
    print(f"语料：{corpus_meta['filename']} · "
          f"选段 {corpus_meta['selected_segments']}/{corpus_meta['full_segments']} · "
          f"{corpus_meta['source_chars']} 字符")

    # 成本预估
    batches = len(core.make_batches(segments))
    per_arm = batches * (1 + 1 + (1 if cfg["run"].get("enable_review") else 0))
    print(f"成本预估：{batches} 批/臂，约 {per_arm} 次 LLM 调用/臂，"
          f"{len(arms)} 臂共约 {per_arm * len(arms)} 次")

    baseline_ref = cfg["code"]["baseline_ref"]
    baseline_wt = ensure_baseline_worktree(baseline_ref)
    current_ref = _git_head(ROOT)
    api_key = os.environ.get("TRANSPRAXIS_EVAL_API_KEY", "")
    tm_seed = Path(cfg["tm_seed"]).resolve() if cfg.get("tm_seed") else None
    if tm_seed is not None and not tm_seed.is_file():
        raise FileNotFoundError(f"TM 种子不存在：{tm_seed}")

    states: Dict[str, Dict[str, Any]] = {}
    for arm in arms:
        code_root = baseline_wt if arm in ("A", "C") else ROOT
        run_dir = out_dir / "runs" / arm
        print(f"\n▶ arm {arm}（code={_git_head(code_root)[:12]}）")
        run_meta = run_arm(
            arm, code_root=code_root, run_dir=run_dir,
            corpus_docx=corpus_docx,
            glossary_file=glossary_file,
            tm_seed=tm_seed if arm in ("C", "D") else None,
            mock=args.mock, run_cfg=cfg["run"], api_key=api_key)
        job_id = run_meta.get("job_id")
        if not job_id:
            raise RuntimeError(f"arm {arm} 缺少 run_meta.job_id")
        state_path = run_dir / job_id / "state.json"
        if not state_path.is_file():
            raise RuntimeError(f"arm {arm} 未生成 state.json：{state_path}")
        states[arm] = json.loads(state_path.read_text(encoding="utf-8"))
        print(f"  完成：segments={run_meta.get('segments')} "
              f"llm_calls={run_meta.get('llm_calls')} "
              f"elapsed={run_meta.get('elapsed_s')}s")

    meta = {
        "baseline_ref": baseline_ref,
        "current_ref": current_ref,
        "corpus": corpus_meta,
        "glossary": str(glossary_file),
        "glossary_auto_locked": args.lock,
        "tm_seed": str(tm_seed) if tm_seed else None,
        "run": cfg["run"],
        "mock": args.mock,
    }
    _finalize(out_dir, states, segments, glossary_entries, meta,
              seed=cfg.get("seed", 42))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
